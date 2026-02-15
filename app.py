import ctypes
from ctypes import wintypes
from datetime import datetime
from difflib import SequenceMatcher
from html.parser import HTMLParser
import io
import json
from pathlib import Path
import re
import threading
import tkinter as tk
import tkinter.font as tkfont
from tkinter import filedialog, messagebox, simpledialog
from xml.sax.saxutils import escape

from PIL import Image, ImageGrab, ImageTk
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

try:
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure
except Exception:
    FigureCanvasAgg = None
    Figure = None

try:
    from wordfreq import zipf_frequency
except Exception:
    zipf_frequency = None


APP_DIR = Path(__file__).resolve().parent
DATA_DIR = APP_DIR / "data"
IMAGES_DIR = DATA_DIR / "images"
SESSION_FILE = DATA_DIR / "session.json"
SESSION_PREV_FILE = DATA_DIR / "session.prev.json"
BACKUP_DIR = DATA_DIR / "backups"
SETTINGS_FILE = DATA_DIR / "settings.json"
CUSTOM_DICTIONARY_FILE = DATA_DIR / "custom_dictionary.json"
AUTOSAVE_MS = 2000
AUTOSAVE_BACKUP_MS = 30000
MAX_SESSION_BACKUPS = 120
CLIPBOARD_POLL_MS = 700
CLIPBOARD_OPEN_RETRIES = 12
CLIPBOARD_OPEN_RETRY_DELAY_MS = 12

IMAGE_TOKEN_PATTERN = re.compile(r"\[\[IMG:([0-9A-Za-z\-]+)\]\]")
MULTISPACE_PATTERN = re.compile(r"[ \t]{2,}")
SPELLCHECK_WORD_PATTERN = re.compile(
    r"[A-Za-z\u00C0-\u024F]+(?:['-][A-Za-z\u00C0-\u024F]+)*"
)
BULLET_MARKERS = ("•", "◦", "▪", "▫")
BULLET_INDENT_SPACES = 4
BULLET_LINE_PATTERN = re.compile(
    r"^([ \t]*)([-*+\u2022\u25E6\u25AA\u25AB])\s"
)
TOC_HEADING_PATTERN = re.compile(r"^\s*\d+(?:\.\d+)*(?:[.)])?\s+\S")
BOLD_TAG = "format_bold"
SUPERSCRIPT_TAG = "format_superscript"
IMAGE_TOKEN_HIDDEN_TAG = "image_token_hidden"
SPELLCHECK_TAG = "spellcheck_misspelled"
IDENTICAL_MATCH_TAG = "identical_match_highlight"
SEARCH_MATCH_TAG = "search_match_highlight"
SEARCH_CURRENT_TAG = "search_current_match"
SPELLCHECK_DEBOUNCE_MS = 350
SPELLCHECK_MIN_ZIPF = 2.6
SPELLCHECK_LANGS = ("en", "nb")
SPELLCHECK_CUSTOM_WORDS = {
    "ctrl",
    "alt",
    "pdf",
    "docx",
    "api",
    "ai",
    "ui",
    "windows",
    "notatapp",
}
INLINE_IMAGE_MAX_WIDTH = 520
INLINE_IMAGE_MAX_HEIGHT = 260
FORMULA_RENDER_DPI = 240
FORMULA_RENDER_FONT_SIZE = 20
BASE_ALPHA = 0.84
WINDOW_GEOMETRY = "680x540+0+0"
DEFAULT_PDF_LINE_SPACING = 1.45
DEFAULT_SETTINGS = {
    "window_alpha": BASE_ALPHA,
    "window_width": 680,
    "window_height": 540,
    "autosave_ms": AUTOSAVE_MS,
    "pdf_line_spacing": DEFAULT_PDF_LINE_SPACING,
}
CLIPBOARD_HTML_FORMAT_NAMES = ("HTML Format", "text/html", "CF_HTML", "HTML")
CLIPBOARD_RTF_FORMAT_NAMES = (
    "Rich Text Format",
    "text/rtf",
    "application/rtf",
    "RTF",
)
CLIPBOARD_STANDARD_FORMATS = {
    1: "CF_TEXT",
    2: "CF_BITMAP",
    3: "CF_METAFILEPICT",
    4: "CF_SYLK",
    5: "CF_DIF",
    6: "CF_TIFF",
    7: "CF_OEMTEXT",
    8: "CF_DIB",
    9: "CF_PALETTE",
    10: "CF_PENDATA",
    11: "CF_RIFF",
    12: "CF_WAVE",
    13: "CF_UNICODETEXT",
    14: "CF_ENHMETAFILE",
    15: "CF_HDROP",
    16: "CF_LOCALE",
    17: "CF_DIBV5",
}


class ClipboardHtmlRunParser(HTMLParser):
    BLOCK_TAGS = {
        "p",
        "div",
        "section",
        "article",
        "header",
        "footer",
        "aside",
        "blockquote",
        "pre",
        "ul",
        "ol",
        "table",
        "tr",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
    }

    def __init__(self, class_bold_map=None, css_vars=None):
        super().__init__(convert_charrefs=True)
        self.runs = []
        self._bold_depth = 0
        self._skip_depth = 0
        self._tag_stack = []
        self._li_depth = 0
        self._class_bold_map = class_bold_map or {}
        self._css_vars = css_vars or {}

    @staticmethod
    def _resolve_css_value(value, css_vars, depth=0):
        if value is None:
            return ""
        if depth > 6:
            return str(value).strip()

        text = str(value).strip()
        if not text:
            return ""

        text = re.sub(r"\s*!important\s*$", "", text, flags=re.IGNORECASE).strip()
        var_match = re.match(r"var\(\s*(--[A-Za-z0-9_-]+)\s*(?:,\s*([^)]+))?\)", text)
        if var_match:
            var_name = var_match.group(1).casefold()
            fallback = var_match.group(2).strip() if var_match.group(2) else ""
            resolved = ""
            if css_vars:
                resolved = css_vars.get(var_name, "")
            if not resolved:
                resolved = fallback
            return ClipboardHtmlRunParser._resolve_css_value(resolved, css_vars, depth + 1)

        return text

    @staticmethod
    def _style_implies_bold(style_text, css_vars=None):
        if not style_text:
            return False

        lowered = style_text.casefold()

        font_match = re.search(r"font\s*:\s*([^;]+)", lowered)
        if font_match:
            font_value = ClipboardHtmlRunParser._resolve_css_value(
                font_match.group(1), css_vars or {}
            )
            if re.search(r"\b(bold|bolder|semibold|demibold|black)\b", font_value):
                return True
            if re.search(r"\b([5-9]00)\b", font_value):
                return True

        match = re.search(r"font-weight\s*:\s*([^;]+)", lowered)
        if match:
            value = ClipboardHtmlRunParser._resolve_css_value(
                match.group(1), css_vars or {}
            ).strip()
            if value in {"bold", "bolder", "semibold", "demibold", "medium", "black"}:
                return True

            number_match = re.match(r"([0-9]{3,4})", value)
            if number_match:
                try:
                    if int(number_match.group(1)) >= 500:
                        return True
                except Exception:
                    pass

        variation_match = re.search(r"font-variation-settings\s*:\s*([^;]+)", lowered)
        if variation_match:
            variation_value = ClipboardHtmlRunParser._resolve_css_value(
                variation_match.group(1), css_vars or {}
            )
            wght_match = re.search(r"['\"]?wght['\"]?\s*([0-9]{2,4})", variation_value)
            if wght_match:
                try:
                    if int(wght_match.group(1)) >= 500:
                        return True
                except Exception:
                    pass

        family_match = re.search(r"font-family\s*:\s*([^;]+)", lowered)
        if family_match:
            family_value = ClipboardHtmlRunParser._resolve_css_value(
                family_match.group(1), css_vars or {}
            )
            if (
                "bold" in family_value
                or "black" in family_value
                or "semibold" in family_value
                or "demibold" in family_value
            ):
                return True

        return False

    def _attrs_imply_bold(self, attrs):
        attr_map = {k.lower(): (v or "") for k, v in attrs}
        if self._style_implies_bold(attr_map.get("style", ""), self._css_vars):
            return True

        face_value = attr_map.get("face", "").casefold()
        if face_value and (
            "bold" in face_value
            or "black" in face_value
            or "semibold" in face_value
            or "demibold" in face_value
        ):
            return True

        class_value = attr_map.get("class", "")
        if class_value:
            for class_name in re.split(r"\s+", class_value.strip()):
                if class_name and self._class_bold_map.get(class_name):
                    return True

        return False

    def _append(self, text, is_bold):
        if not text:
            return
        if self.runs and self.runs[-1][1] == is_bold:
            prev_text, _ = self.runs[-1]
            self.runs[-1] = (prev_text + text, is_bold)
            return
        self.runs.append((text, is_bold))

    def _append_newline(self):
        if not self.runs:
            return
        if self.runs[-1][0].endswith("\n"):
            return
        self._append("\n", False)

    def handle_starttag(self, tag, attrs):
        lowered = tag.lower()
        if lowered in ("script", "style"):
            self._skip_depth += 1
            return

        if self._skip_depth > 0:
            return

        if lowered == "br":
            self._append("\n", False)
            return

        if lowered == "li":
            self._append_newline()
            self._append("- ", False)
            self._li_depth += 1

        pushes_bold = lowered in ("b", "strong", "th")
        if len(lowered) == 2 and lowered.startswith("h") and lowered[1].isdigit():
            pushes_bold = True
        if self._attrs_imply_bold(attrs):
            pushes_bold = True

        if pushes_bold:
            self._bold_depth += 1

        self._tag_stack.append((lowered, pushes_bold))

    def handle_endtag(self, tag):
        lowered = tag.lower()
        if lowered in ("script", "style"):
            if self._skip_depth > 0:
                self._skip_depth -= 1
            return

        if self._skip_depth > 0:
            return

        for i in range(len(self._tag_stack) - 1, -1, -1):
            stack_tag, pushes_bold = self._tag_stack[i]
            if stack_tag != lowered:
                continue
            self._tag_stack.pop(i)
            if pushes_bold and self._bold_depth > 0:
                self._bold_depth -= 1
            break

        if lowered == "li":
            if self._li_depth > 0:
                self._li_depth -= 1
            self._append_newline()
        elif lowered in self.BLOCK_TAGS:
            if self._li_depth == 0:
                self._append_newline()

    def handle_data(self, data):
        if self._skip_depth > 0:
            return
        self._append(data, self._bold_depth > 0)


class ClipboardRtfRunParser:
    def __init__(self):
        self.runs = []
        self._bold = False
        self._bold_stack = []
        self._uc_skip = 1
        self._pending_skip = 0

    def _append(self, text, is_bold=None):
        if not text:
            return
        if is_bold is None:
            is_bold = self._bold
        if self.runs and self.runs[-1][1] == is_bold:
            prev_text, _ = self.runs[-1]
            self.runs[-1] = (prev_text + text, is_bold)
            return
        self.runs.append((text, is_bold))

    def _read_control_word(self, text, index):
        n = len(text)
        start = index
        while index < n and text[index].isalpha():
            index += 1
        word = text[start:index]

        sign = 1
        if index < n and text[index] in ("-", "+"):
            sign = -1 if text[index] == "-" else 1
            index += 1

        number_start = index
        while index < n and text[index].isdigit():
            index += 1
        number = None
        if number_start < index:
            number = sign * int(text[number_start:index])

        if index < n and text[index] == " ":
            index += 1

        return word, number, index

    def parse(self, rtf_text):
        self.runs = []
        self._bold = False
        self._bold_stack = []
        self._uc_skip = 1
        self._pending_skip = 0

        i = 0
        n = len(rtf_text)
        while i < n:
            ch = rtf_text[i]

            if self._pending_skip > 0 and ch not in ("\\", "{", "}"):
                self._pending_skip -= 1
                i += 1
                continue

            if ch == "{":
                self._bold_stack.append(self._bold)
                i += 1
                continue

            if ch == "}":
                if self._bold_stack:
                    self._bold = self._bold_stack.pop()
                i += 1
                continue

            if ch != "\\":
                self._append(ch)
                i += 1
                continue

            i += 1
            if i >= n:
                break

            symbol = rtf_text[i]
            if symbol in ("\\", "{", "}"):
                self._append(symbol)
                i += 1
                continue

            if symbol == "'":
                if i + 2 < n:
                    hex_value = rtf_text[i + 1 : i + 3]
                    try:
                        decoded = bytes.fromhex(hex_value).decode("cp1252")
                        self._append(decoded)
                    except Exception:
                        pass
                    i += 3
                else:
                    i += 1
                continue

            if not symbol.isalpha():
                if symbol == "~":
                    self._append(" ")
                elif symbol == "_":
                    self._append("-")
                elif symbol == "-":
                    self._append("-")
                elif symbol == "*":
                    pass
                i += 1
                continue

            word, number, i = self._read_control_word(rtf_text, i)
            if word == "b":
                self._bold = False if number == 0 else True
            elif word in ("par", "line"):
                self._append("\n", False)
            elif word == "tab":
                self._append("\t")
            elif word in ("emdash", "endash"):
                self._append("-")
            elif word == "bullet":
                self._append("- ", False)
            elif word == "u" and number is not None:
                codepoint = number if number >= 0 else number + 65536
                try:
                    self._append(chr(codepoint))
                except Exception:
                    pass
                self._pending_skip = max(0, self._uc_skip)
            elif word == "uc" and number is not None and number >= 0:
                self._uc_skip = number

        return self.runs


class GlobalHotkeyListener:
    WM_HOTKEY = 0x0312
    WM_QUIT = 0x0012
    MOD_ALT = 0x0001
    MOD_CONTROL = 0x0002
    MOD_SHIFT = 0x0004
    HOTKEY_CHAR = "|"

    def __init__(self, on_hotkey, on_error=None):
        self._on_hotkey = on_hotkey
        self._on_error = on_error
        self._thread = threading.Thread(target=self._run, daemon=True)
        self._thread_id = None
        self._registered = False

    def start(self):
        self._thread.start()

    def stop(self):
        if self._thread_id is None:
            return
        user32 = ctypes.windll.user32
        user32.PostThreadMessageW(self._thread_id, self.WM_QUIT, 0, 0)
        self._thread.join(timeout=1.0)

    def _run(self):
        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32
        self._thread_id = kernel32.GetCurrentThreadId()

        vk_data = user32.VkKeyScanW(ord(self.HOTKEY_CHAR))
        if vk_data == -1:
            if self._on_error is not None:
                self._on_error("Kunne ikke mappe hurtigtast '|'.")
            return

        vk = vk_data & 0xFF
        shift_state = (vk_data >> 8) & 0xFF
        modifiers = 0
        if shift_state & 0x01:
            modifiers |= self.MOD_SHIFT
        if shift_state & 0x02:
            modifiers |= self.MOD_CONTROL
        if shift_state & 0x04:
            modifiers |= self.MOD_ALT

        ok = user32.RegisterHotKey(None, 1, modifiers, vk)
        if not ok:
            if self._on_error is not None:
                self._on_error(
                    "Kunne ikke registrere hurtigtast '|'. Prøv å lukke andre apper som bruker samme hurtigtast."
                )
            return

        self._registered = True
        msg = wintypes.MSG()
        while True:
            result = user32.GetMessageW(ctypes.byref(msg), None, 0, 0)
            if result in (0, -1):
                break
            if msg.message == self.WM_HOTKEY and msg.wParam == 1:
                self._on_hotkey()

        if self._registered:
            user32.UnregisterHotKey(None, 1)
            self._registered = False


class NoteOverlayApp:
    def __init__(self):
        DATA_DIR.mkdir(exist_ok=True)
        IMAGES_DIR.mkdir(exist_ok=True)
        BACKUP_DIR.mkdir(exist_ok=True)
        self.settings = self._load_settings()
        self._user_spell_words = self._load_user_dictionary_words()
        self.custom_spell_words = set(SPELLCHECK_CUSTOM_WORDS)
        self.custom_spell_words.update(self._user_spell_words)
        self.autosave_ms = int(self.settings.get("autosave_ms", AUTOSAVE_MS))
        self.pdf_line_spacing = float(
            self.settings.get("pdf_line_spacing", DEFAULT_PDF_LINE_SPACING)
        )
        self._init_clipboard_api()

        


        self.root = tk.Tk()
        self.root.title("Notat Overlay")
        self.root.geometry(self._geometry_from_settings())
        self.root.minsize(620, 420)
        self.root.attributes("-topmost", True)
        self.root.attributes("-alpha", self._target_alpha())
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        self.theme = {
            "bg": "#ecf1f6",
            "card": "#fbfcfd",
            "card_soft": "#f3f6f9",
            "border": "#d5dde8",
            "text": "#1d2733",
            "muted": "#64758b",
            "accent": "#2d6cdf",
            "accent_hover": "#2158bb",
            "button": "#e5ebf4",
            "button_hover": "#d9e2ef",
            "status": "#e7edf5",
        }
        self.ui_font = self._pick_font(("Segoe UI Variable Text", "Segoe UI", "Calibri"), 10)
        self.title_font = self._pick_font(("Segoe UI Semibold", "Segoe UI", "Calibri"), 12)
        self.text_font = self._pick_font(("Segoe UI", "Calibri"), 11)
        self.text_bold_font = (self.text_font[0], self.text_font[1], "bold")

        self.root.configure(bg=self.theme["bg"])
        self._fade_after_id = None

        self.attachments = {}
        self.formula_meta = {}
        self._dirty = False
        self.auto_capture_enabled = False
        self._last_clipboard_signature = None
        self._inline_image_refs = []
        self.bold_typing_mode = False
        self._spellcheck_after_id = None
        self.spellcheck_enabled = True
        self._spellcheck_available = zipf_frequency is not None
        self._last_backup_snapshot_key = None
        self._last_backup_snapshot_at = None
        self._last_identical_token = ""
        self.focus_mode = False
        self.search_panel_visible = False
        self._search_match_ranges = []
        self._search_current_idx = -1
        self.search_var = tk.StringVar()
        self.replace_var = tk.StringVar()
        self.search_case_var = tk.BooleanVar(value=False)
        self._toc_window = None
        self._toc_listbox = None
        self._toc_entries = []
        self._settings_window = None
        self._context_menu = None

        self.status_var = tk.StringVar(
            value="'|': vis/skjul. Auto-fangst: marker tekst + Ctrl+C."
        )
        self._build_ui()
        self._refresh_bold_button()
        self._refresh_spellcheck_button()
        self._refresh_focus_button()
        self._load_session()

        self.hotkey = GlobalHotkeyListener(
            on_hotkey=lambda: self.root.after(0, self.toggle_visibility),
            on_error=lambda msg: self.root.after(0, self._set_status, msg),
        )
        self.hotkey.start()

        self.root.after(self.autosave_ms, self._autosave_tick)
        self.root.after(CLIPBOARD_POLL_MS, self._clipboard_watch_tick)
        self._schedule_spellcheck()
        self.text.focus_set()
    def _open_clipboard_with_retry(self):
            user32 = self._u32
            for _ in range(CLIPBOARD_OPEN_RETRIES):
                if user32.OpenClipboard(None):
                    return True
                try:
                    self.root.update_idletasks()
                    self.root.after(CLIPBOARD_OPEN_RETRY_DELAY_MS)
                except Exception:
                    pass
            return False

    def _close_clipboard_quietly(self):
        try:
            self._u32.CloseClipboard()
        except Exception:
            pass

    def _init_clipboard_api(self):
        self._u32 = ctypes.WinDLL("user32", use_last_error=True)
        self._k32 = ctypes.WinDLL("kernel32", use_last_error=True)

        self._u32.RegisterClipboardFormatW.argtypes = [wintypes.LPCWSTR]
        self._u32.RegisterClipboardFormatW.restype = wintypes.UINT
        self._u32.OpenClipboard.argtypes = [wintypes.HWND]
        self._u32.OpenClipboard.restype = wintypes.BOOL
        self._u32.CloseClipboard.argtypes = []
        self._u32.CloseClipboard.restype = wintypes.BOOL
        self._u32.IsClipboardFormatAvailable.argtypes = [wintypes.UINT]
        self._u32.IsClipboardFormatAvailable.restype = wintypes.BOOL
        self._u32.GetClipboardData.argtypes = [wintypes.UINT]
        self._u32.GetClipboardData.restype = wintypes.HANDLE
        self._u32.EnumClipboardFormats.argtypes = [wintypes.UINT]
        self._u32.EnumClipboardFormats.restype = wintypes.UINT
        self._u32.GetClipboardFormatNameW.argtypes = [
            wintypes.UINT,
            wintypes.LPWSTR,
            ctypes.c_int,
        ]
        self._u32.GetClipboardFormatNameW.restype = ctypes.c_int

        self._k32.GlobalSize.argtypes = [ctypes.c_void_p]
        self._k32.GlobalSize.restype = ctypes.c_size_t
        self._k32.GlobalLock.argtypes = [ctypes.c_void_p]
        self._k32.GlobalLock.restype = ctypes.c_void_p
        self._k32.GlobalUnlock.argtypes = [ctypes.c_void_p]
        self._k32.GlobalUnlock.restype = wintypes.BOOL

    def _pick_font(self, candidates, size):
        available = set(tkfont.families())
        for candidate in candidates:
            if candidate in available:
                return (candidate, size)
        return ("TkDefaultFont", size)

    def _sanitize_settings(self, payload):
        merged = dict(DEFAULT_SETTINGS)
        if not isinstance(payload, dict):
            return merged

        def as_float(value, fallback):
            try:
                return float(value)
            except Exception:
                return fallback

        def as_int(value, fallback):
            try:
                return int(value)
            except Exception:
                return fallback

        merged["window_alpha"] = min(
            0.98, max(0.55, as_float(payload.get("window_alpha"), merged["window_alpha"]))
        )
        merged["window_width"] = min(
            1200, max(560, as_int(payload.get("window_width"), merged["window_width"]))
        )
        merged["window_height"] = min(
            980, max(420, as_int(payload.get("window_height"), merged["window_height"]))
        )
        merged["autosave_ms"] = min(
            60000, max(1000, as_int(payload.get("autosave_ms"), merged["autosave_ms"]))
        )
        merged["pdf_line_spacing"] = min(
            2.4,
            max(1.15, as_float(payload.get("pdf_line_spacing"), merged["pdf_line_spacing"])),
        )
        return merged

    def _load_settings(self):
        if not SETTINGS_FILE.exists():
            return dict(DEFAULT_SETTINGS)
        try:
            payload = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        except Exception:
            return dict(DEFAULT_SETTINGS)
        return self._sanitize_settings(payload)

    def _save_settings(self):
        clean = self._sanitize_settings(self.settings)
        self.settings = clean
        SETTINGS_FILE.write_text(
            json.dumps(clean, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    def _geometry_from_settings(self):
        width = int(self.settings.get("window_width", DEFAULT_SETTINGS["window_width"]))
        height = int(self.settings.get("window_height", DEFAULT_SETTINGS["window_height"]))
        return f"{width}x{height}+0+0"

    def _target_alpha(self):
        try:
            return float(self.settings.get("window_alpha", BASE_ALPHA))
        except Exception:
            return BASE_ALPHA

    def _load_user_dictionary_words(self):
        if not CUSTOM_DICTIONARY_FILE.exists():
            return set()
        try:
            payload = json.loads(CUSTOM_DICTIONARY_FILE.read_text(encoding="utf-8"))
        except Exception:
            return set()
        if not isinstance(payload, list):
            return set()
        words = set()
        for item in payload:
            if not isinstance(item, str):
                continue
            normalized = item.casefold().strip()
            if normalized:
                words.add(normalized)
        return words

    def _save_user_dictionary_words(self):
        payload = sorted(self._user_spell_words)
        CUSTOM_DICTIONARY_FILE.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    def _create_button(self, parent, text, command, primary=False):
        bg = self.theme["accent"] if primary else self.theme["button"]
        fg = "#ffffff" if primary else self.theme["text"]
        hover = self.theme["accent_hover"] if primary else self.theme["button_hover"]
        button = tk.Button(
            parent,
            text=text,
            command=command,
            bg=bg,
            fg=fg,
            activebackground=hover,
            activeforeground=fg,
            borderwidth=0,
            relief="flat",
            padx=12,
            pady=6,
            font=self.ui_font,
            cursor="hand2",
            highlightthickness=0,
        )
        button.bind("<Enter>", lambda _event: button.configure(bg=hover))
        button.bind("<Leave>", lambda _event: button.configure(bg=bg))
        return button

    def _fade_to(self, target, steps=10, interval_ms=14, on_complete=None):
        if self._fade_after_id is not None:
            self.root.after_cancel(self._fade_after_id)
            self._fade_after_id = None

        current = float(self.root.attributes("-alpha"))
        delta = (target - current) / max(steps, 1)

        def animate(step_index=0, value=current):
            if step_index >= steps:
                self.root.attributes("-alpha", target)
                self._fade_after_id = None
                if on_complete is not None:
                    on_complete()
                return

            value += delta
            self.root.attributes("-alpha", min(1.0, max(0.0, value)))
            self._fade_after_id = self.root.after(
                interval_ms, lambda: animate(step_index + 1, value)
            )

        animate()

    def _build_ui(self):
        self.container_frame = tk.Frame(self.root, bg=self.theme["bg"])
        self.container_frame.pack(fill="both", expand=True, padx=14, pady=14)

        self.header_frame = tk.Frame(
            self.container_frame,
            bg=self.theme["card"],
            highlightthickness=1,
            highlightbackground=self.theme["border"],
            highlightcolor=self.theme["border"],
        )
        self.header_frame.pack(fill="x")

        tk.Label(
            self.header_frame,
            text="Quick Notes",
            bg=self.theme["card"],
            fg=self.theme["text"],
            font=self.title_font,
            padx=12,
            pady=8,
        ).pack(side="left")

        tk.Label(
            self.header_frame,
            text="|",
            bg=self.theme["card_soft"],
            fg=self.theme["muted"],
            font=self.ui_font,
            padx=10,
            pady=5,
        ).pack(side="right", padx=10, pady=8)

        self.toolbar_frame = tk.Frame(self.container_frame, bg=self.theme["bg"])
        self.toolbar_frame.pack(fill="x", pady=(10, 8))

        btn_new = self._create_button(self.toolbar_frame, "Nytt", self.new_note)
        btn_new.pack(side="left")

        btn_word = self._create_button(self.toolbar_frame, "Word", self.export_word)
        btn_word.pack(side="left", padx=6)

        btn_pdf = self._create_button(self.toolbar_frame, "PDF", self.export_pdf)
        btn_pdf.pack(side="left")

        self.btn_auto_capture = self._create_button(
            self.toolbar_frame, "Auto: AV", self.toggle_auto_capture
        )
        self.btn_auto_capture.pack(side="left", padx=(6, 0))

        btn_search = self._create_button(
            self.toolbar_frame, "Søk", lambda: self.open_search_panel(replace_mode=False)
        )
        btn_search.pack(side="left", padx=(6, 0))

        btn_formula = self._create_button(self.toolbar_frame, "f(x)", self.insert_formula)
        btn_formula.pack(side="left", padx=(6, 0))

        self.btn_bold = self._create_button(self.toolbar_frame, "B", self.toggle_bold)

        btn_bullet = self._create_button(self.toolbar_frame, "Bullet", self.toggle_bullets)

        self.btn_spellcheck = self._create_button(
            self.toolbar_frame, "Stavekontroll: PÅ", self.toggle_spellcheck
        )

        btn_mark_identical = self._create_button(
            self.toolbar_frame, "Marker like", self.mark_identical
        )

        btn_replace_identical = self._create_button(
            self.toolbar_frame, "Erstatt like", self.replace_identical
        )

        btn_toc = self._create_button(self.toolbar_frame, "Innhold", self.show_toc)
        btn_toc.pack(side="right", padx=(0, 6))

        btn_settings = self._create_button(self.toolbar_frame, "Innst.", self.show_settings)
        btn_settings.pack(side="right", padx=(0, 6))

        btn_help = self._create_button(self.toolbar_frame, "Hjelp", self.show_help)
        btn_help.pack(side="right", padx=(0, 6))

        self.btn_focus_mode = self._create_button(
            self.toolbar_frame, "Fokus: AV", self.toggle_focus_mode
        )
        self.btn_focus_mode.pack(side="right", padx=(0, 6))

        btn_hide = self._create_button(
            self.toolbar_frame, "Skjul", self.toggle_visibility, primary=True
        )
        btn_hide.pack(side="right", padx=(0, 6))

        self.search_panel = tk.Frame(
            self.container_frame,
            bg=self.theme["card_soft"],
            highlightthickness=1,
            highlightbackground=self.theme["border"],
            highlightcolor=self.theme["border"],
        )

        tk.Label(
            self.search_panel,
            text="Finn",
            bg=self.theme["card_soft"],
            fg=self.theme["text"],
            font=self.ui_font,
            padx=8,
        ).pack(side="left")
        self.entry_search = tk.Entry(
            self.search_panel,
            textvariable=self.search_var,
            relief="flat",
            borderwidth=0,
            bg=self.theme["card"],
            fg=self.theme["text"],
            insertbackground=self.theme["text"],
            width=24,
        )
        self.entry_search.pack(side="left", padx=(0, 6), pady=6, ipady=4)
        self.entry_search.bind("<KeyRelease>", self._on_search_query_changed)
        self.entry_search.bind("<Return>", self._on_search_next_shortcut)
        self.entry_search.bind("<Shift-Return>", self._on_search_prev_shortcut)

        self.chk_search_case = tk.Checkbutton(
            self.search_panel,
            text="Aa",
            variable=self.search_case_var,
            command=self._refresh_search_matches,
            bg=self.theme["card_soft"],
            fg=self.theme["text"],
            selectcolor=self.theme["card_soft"],
            activebackground=self.theme["card_soft"],
            activeforeground=self.theme["text"],
            font=self.ui_font,
            padx=4,
        )
        self.chk_search_case.pack(side="left")

        btn_prev = self._create_button(self.search_panel, "Forrige", self.search_prev)
        btn_prev.pack(side="left", padx=(4, 2))

        btn_next = self._create_button(self.search_panel, "Neste", self.search_next)
        btn_next.pack(side="left", padx=(0, 8))

        tk.Label(
            self.search_panel,
            text="Erstatt",
            bg=self.theme["card_soft"],
            fg=self.theme["text"],
            font=self.ui_font,
            padx=4,
        ).pack(side="left")

        self.entry_replace = tk.Entry(
            self.search_panel,
            textvariable=self.replace_var,
            relief="flat",
            borderwidth=0,
            bg=self.theme["card"],
            fg=self.theme["text"],
            insertbackground=self.theme["text"],
            width=18,
        )
        self.entry_replace.pack(side="left", padx=(2, 6), pady=6, ipady=4)
        self.entry_replace.bind("<Return>", self._on_replace_one_shortcut)

        btn_replace_one = self._create_button(
            self.search_panel, "Erstatt", self.replace_current_match
        )
        btn_replace_one.pack(side="left", padx=(0, 2))

        btn_replace_all = self._create_button(
            self.search_panel, "Alle", self.replace_all_matches
        )
        btn_replace_all.pack(side="left")

        btn_close_search = self._create_button(
            self.search_panel, "Lukk", self.hide_search_panel
        )
        btn_close_search.pack(side="right", padx=6)

        self.text_frame = tk.Frame(
            self.container_frame,
            bg=self.theme["card"],
            highlightthickness=1,
            highlightbackground=self.theme["border"],
            highlightcolor=self.theme["border"],
        )
        self.text_frame.pack(fill="both", expand=True)
        self.text_frame.grid_rowconfigure(0, weight=1)
        self.text_frame.grid_columnconfigure(0, weight=1)

        self.text = tk.Text(
            self.text_frame,
            wrap="word",
            undo=True,
            font=self.text_font,
            relief="flat",
            borderwidth=0,
            padx=12,
            pady=12,
            bg=self.theme["card"],
            fg=self.theme["text"],
            insertbackground=self.theme["text"],
            selectbackground="#d8e6ff",
            highlightthickness=0,
            spacing1=2,
            spacing2=4,
            spacing3=2,
        )
        self.text.grid(row=0, column=0, sticky="nsew")
        self.text.tag_configure(BOLD_TAG, font=self.text_bold_font)
        self.text.tag_configure(SUPERSCRIPT_TAG, offset=4)
        self.text.tag_configure(IMAGE_TOKEN_HIDDEN_TAG, elide=True)
        self.text.tag_configure(SPELLCHECK_TAG, underline=True, foreground="#c62828")
        self.text.tag_configure(IDENTICAL_MATCH_TAG, background="#fff1a8")
        self.text.tag_configure(SEARCH_MATCH_TAG, background="#ffe9a8")
        self.text.tag_configure(SEARCH_CURRENT_TAG, background="#ffc36a")

        scroll = tk.Scrollbar(
            self.text_frame,
            command=self.text.yview,
            borderwidth=0,
            relief="flat",
            troughcolor=self.theme["card"],
            bg=self.theme["button"],
            activebackground=self.theme["button_hover"],
        )
        scroll.grid(row=0, column=1, sticky="ns", padx=(0, 2), pady=2)
        self.text.configure(yscrollcommand=scroll.set)

        self.status_label = tk.Label(
            self.container_frame,
            textvariable=self.status_var,
            anchor="w",
            bg=self.theme["status"],
            fg=self.theme["muted"],
            padx=10,
            pady=6,
            font=self.ui_font,
        )
        self.status_label.pack(fill="x", pady=(8, 0))

        self.text.bind("<<Modified>>", self._on_text_modified)
        self.text.bind("<KeyPress>", self._on_text_keypress)
        self.text.bind("<KeyRelease>", self._on_text_keyrelease)
        self.text.bind("<Control-v>", self._on_paste)
        self.text.bind("<Control-V>", self._on_paste)
        self.text.bind("<Control-b>", self._on_bold_shortcut)
        self.text.bind("<Control-B>", self._on_bold_shortcut)
        self.text.bind("<Control-BackSpace>", self._on_ctrl_backspace)
        self.text.bind("<Control-Shift-B>", self._on_bullet_shortcut)
        self.text.bind("<Control-Shift-b>", self._on_bullet_shortcut)
        self.text.bind("<Control-Shift-M>", self._on_mark_identical_shortcut)
        self.text.bind("<Control-Shift-m>", self._on_mark_identical_shortcut)
        self.text.bind("<Control-Shift-R>", self._on_replace_identical_shortcut)
        self.text.bind("<Control-Shift-r>", self._on_replace_identical_shortcut)
        self.text.bind("<Control-Shift-L>", self._on_spellcheck_shortcut)
        self.text.bind("<Control-Shift-l>", self._on_spellcheck_shortcut)
        self.text.bind("<Control-f>", self._on_find_shortcut)
        self.text.bind("<Control-h>", self._on_replace_shortcut)
        self.text.bind("<Control-m>", self._on_formula_shortcut)
        self.text.bind("<Control-M>", self._on_formula_shortcut)
        self.text.bind("<Control-comma>", self._on_settings_shortcut)
        self.text.bind("<F1>", self._on_help_shortcut)
        self.text.bind("<Button-3>", self._on_text_context_menu)
        self.text.bind("<Tab>", self._on_tab_indent)
        self.text.bind("<Shift-Tab>", self._on_shift_tab_outdent)
        self.text.bind("<ISO_Left_Tab>", self._on_shift_tab_outdent)
        self.text.bind("<Return>", self._on_return)
        self.root.bind("<Control-s>", self._on_ctrl_s)
        self.root.bind("<Control-f>", self._on_find_shortcut)
        self.root.bind("<Control-h>", self._on_replace_shortcut)
        self.root.bind("<Control-m>", self._on_formula_shortcut)
        self.root.bind("<Control-M>", self._on_formula_shortcut)
        self.root.bind("<Control-Shift-f>", self._on_focus_mode_shortcut)
        self.root.bind("<Control-Shift-F>", self._on_focus_mode_shortcut)
        self.root.bind("<Control-Shift-t>", self._on_toc_shortcut)
        self.root.bind("<Control-Shift-T>", self._on_toc_shortcut)
        self.root.bind("<Control-Shift-L>", self._on_spellcheck_shortcut)
        self.root.bind("<Control-Shift-l>", self._on_spellcheck_shortcut)
        self.root.bind("<Control-comma>", self._on_settings_shortcut)
        self.root.bind("<F1>", self._on_help_shortcut)
        self.root.bind("<Escape>", self._on_escape)

    def _on_ctrl_s(self, _event):
        self._save_session(silent=False)
        return "break"

    def _on_escape(self, _event):
        if self.search_panel_visible:
            self.hide_search_panel()
            return "break"
        self.toggle_visibility()
        return "break"

    def _on_settings_shortcut(self, _event):
        self.show_settings()
        return "break"

    def _on_help_shortcut(self, _event):
        self.show_help()
        return "break"

    def _on_formula_shortcut(self, _event):
        self.insert_formula()
        return "break"

    def show_help(self):
        shortcuts = [
            "Ctrl+S: lagre nå",
            "|: vis/skjul notatvindu",
            "Ctrl+F / Ctrl+H: søk / erstatt",
            "Ctrl+M: sett inn matteformel (LaTeX)",
            "Ctrl+Shift+F: fokusmodus",
            "Ctrl+Shift+T: innholdsfortegnelse",
            "Ctrl+B: bold på/av",
            "Ctrl+Shift+B: bullet på/av",
            "Tab / Shift+Tab: bullet indent ut/inn",
            "Ctrl+Backspace: slett forrige ord",
            "Ctrl+Shift+M / Ctrl+Shift+R: marker like / erstatt like",
            "Ctrl+Shift+L: stavekontroll på/av",
            "Ctrl+, : innstillinger",
            "F1: denne hjelpen",
            "Høyreklikk på ord: legg til i ordbok",
        ]
        messagebox.showinfo(
            "Hjelp og Shortcuts",
            "Snarveier:\n\n" + "\n".join(shortcuts),
            parent=self.root,
        )

    def show_settings(self):
        if self._settings_window is not None and self._settings_window.winfo_exists():
            self._settings_window.deiconify()
            self._settings_window.lift()
            self._settings_window.focus_force()
            return

        win = tk.Toplevel(self.root)
        self._settings_window = win
        win.title("Innstillinger")
        win.geometry("420x320+740+90")
        win.attributes("-topmost", True)
        win.configure(bg=self.theme["bg"])

        panel = tk.Frame(
            win,
            bg=self.theme["card"],
            highlightthickness=1,
            highlightbackground=self.theme["border"],
            highlightcolor=self.theme["border"],
        )
        panel.pack(fill="both", expand=True, padx=10, pady=10)

        alpha_var = tk.DoubleVar(value=round(self._target_alpha() * 100))
        autosave_var = tk.IntVar(value=max(1, int(round(self.autosave_ms / 1000))))
        width_var = tk.IntVar(value=int(self.settings.get("window_width", 680)))
        height_var = tk.IntVar(value=int(self.settings.get("window_height", 540)))
        pdf_spacing_var = tk.DoubleVar(value=float(self.pdf_line_spacing))

        tk.Label(
            panel, text="Transparens (%)", bg=self.theme["card"], fg=self.theme["text"], font=self.ui_font
        ).grid(row=0, column=0, sticky="w", padx=12, pady=(12, 4))
        tk.Scale(
            panel,
            from_=55,
            to=98,
            orient="horizontal",
            variable=alpha_var,
            bg=self.theme["card"],
            fg=self.theme["text"],
            troughcolor=self.theme["card_soft"],
            highlightthickness=0,
            length=240,
        ).grid(row=1, column=0, columnspan=2, sticky="w", padx=10)

        tk.Label(
            panel, text="Autosave (sek)", bg=self.theme["card"], fg=self.theme["text"], font=self.ui_font
        ).grid(row=2, column=0, sticky="w", padx=12, pady=(10, 2))
        tk.Spinbox(
            panel,
            from_=1,
            to=60,
            textvariable=autosave_var,
            width=6,
            relief="flat",
            borderwidth=1,
            bg=self.theme["card_soft"],
        ).grid(row=2, column=1, sticky="w", padx=6)

        tk.Label(
            panel, text="Vindusbredde", bg=self.theme["card"], fg=self.theme["text"], font=self.ui_font
        ).grid(row=3, column=0, sticky="w", padx=12, pady=(8, 2))
        tk.Spinbox(
            panel,
            from_=560,
            to=1200,
            increment=20,
            textvariable=width_var,
            width=8,
            relief="flat",
            borderwidth=1,
            bg=self.theme["card_soft"],
        ).grid(row=3, column=1, sticky="w", padx=6)

        tk.Label(
            panel, text="Vindushøyde", bg=self.theme["card"], fg=self.theme["text"], font=self.ui_font
        ).grid(row=4, column=0, sticky="w", padx=12, pady=(8, 2))
        tk.Spinbox(
            panel,
            from_=420,
            to=980,
            increment=20,
            textvariable=height_var,
            width=8,
            relief="flat",
            borderwidth=1,
            bg=self.theme["card_soft"],
        ).grid(row=4, column=1, sticky="w", padx=6)

        tk.Label(
            panel,
            text="PDF linjeavstand",
            bg=self.theme["card"],
            fg=self.theme["text"],
            font=self.ui_font,
        ).grid(row=5, column=0, sticky="w", padx=12, pady=(8, 2))
        tk.Scale(
            panel,
            from_=1.15,
            to=2.40,
            resolution=0.05,
            orient="horizontal",
            variable=pdf_spacing_var,
            bg=self.theme["card"],
            fg=self.theme["text"],
            troughcolor=self.theme["card_soft"],
            highlightthickness=0,
            length=180,
        ).grid(row=5, column=1, sticky="w", padx=6)

        def apply_settings():
            updated = {
                "window_alpha": float(alpha_var.get()) / 100.0,
                "autosave_ms": int(autosave_var.get()) * 1000,
                "window_width": int(width_var.get()),
                "window_height": int(height_var.get()),
                "pdf_line_spacing": float(pdf_spacing_var.get()),
            }
            self.settings.update(updated)
            self.settings = self._sanitize_settings(self.settings)
            self.autosave_ms = int(self.settings["autosave_ms"])
            self.pdf_line_spacing = float(self.settings["pdf_line_spacing"])
            self.root.attributes("-alpha", self._target_alpha())
            self.root.geometry(self._geometry_from_settings())
            self._save_settings()
            self._set_status("Innstillinger lagret.")

        def close_window():
            self._settings_window = None
            win.destroy()

        button_row = tk.Frame(panel, bg=self.theme["card"])
        button_row.grid(row=6, column=0, columnspan=2, sticky="e", padx=10, pady=(14, 10))
        self._create_button(button_row, "Lagre", apply_settings).pack(side="right")
        self._create_button(button_row, "Lukk", close_window).pack(side="right", padx=(0, 6))

        win.protocol("WM_DELETE_WINDOW", close_window)

    def _word_at_index(self, index):
        line_start = self.text.index(f"{index} linestart")
        line_end = self.text.index(f"{index} lineend")
        line_text = self.text.get(line_start, line_end)
        column = int(index.split(".")[1])

        for match in SPELLCHECK_WORD_PATTERN.finditer(line_text):
            if match.start() <= column <= match.end():
                start = f"{line_start}+{match.start()}c"
                end = f"{line_start}+{match.end()}c"
                return match.group(0), start, end
        return "", "", ""

    def _add_word_to_dictionary(self, word):
        normalized = word.casefold().strip()
        if len(normalized) <= 1:
            self._set_status("Kunne ikke legge til ord i ordbok.")
            return
        if normalized in self.custom_spell_words:
            self._set_status(f"'{word}' finnes allerede i ordbok.")
            return

        self._user_spell_words.add(normalized)
        self.custom_spell_words.add(normalized)
        self._save_user_dictionary_words()
        self._schedule_spellcheck()
        self._set_status(f"La til '{word}' i personlig ordbok.")

    def _on_text_context_menu(self, event):
        index = self.text.index(f"@{event.x},{event.y}")
        self.text.mark_set("insert", index)

        word, word_start, word_end = self._word_at_index(index)
        if word_start and word_end:
            self.text.tag_remove("sel", "1.0", "end")
            self.text.tag_add("sel", word_start, word_end)

        menu = tk.Menu(self.root, tearoff=0)
        if word:
            normalized = word.casefold().strip()
            if normalized in self.custom_spell_words:
                menu.add_command(label=f"I ordbok: {word}", state="disabled")
            else:
                menu.add_command(
                    label=f"Legg til i ordbok: {word}",
                    command=lambda w=word: self._add_word_to_dictionary(w),
                )
        else:
            menu.add_command(label="Legg til i ordbok", state="disabled")

        menu.add_separator()
        menu.add_command(label="Sett inn formel", command=self.insert_formula)
        menu.add_separator()
        menu.add_command(label="Søk", command=lambda: self.open_search_panel(False))
        menu.add_command(label="Erstatt like", command=self.replace_identical)
        menu.add_command(label="Marker like", command=self.mark_identical)

        self._context_menu = menu
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

        return "break"

    def _refresh_focus_button(self):
        label = "Fokus: PÅ" if self.focus_mode else "Fokus: AV"
        self.btn_focus_mode.configure(text=label)

    def _set_focus_layout(self):
        if self.focus_mode:
            self.header_frame.pack_forget()
            self.toolbar_frame.pack_forget()
            self.container_frame.pack_configure(padx=8, pady=8)
            self.status_label.configure(pady=4)
        else:
            if not self.header_frame.winfo_manager():
                before_widget = (
                    self.toolbar_frame if self.toolbar_frame.winfo_manager() else self.text_frame
                )
                self.header_frame.pack(fill="x", before=before_widget)
            if not self.toolbar_frame.winfo_manager():
                self.toolbar_frame.pack(fill="x", pady=(10, 8), before=self.text_frame)
            self.container_frame.pack_configure(padx=14, pady=14)
            self.status_label.configure(pady=6)

    def toggle_focus_mode(self):
        self.focus_mode = not self.focus_mode
        self._set_focus_layout()
        self._refresh_focus_button()
        if self.focus_mode:
            self._set_status("Fokusmodus PÅ. Ctrl+Shift+F for å avslutte.")
        else:
            self._set_status("Fokusmodus AV.")

    def _on_focus_mode_shortcut(self, _event):
        self.toggle_focus_mode()
        return "break"

    def _set_search_panel_visible(self, visible):
        self.search_panel_visible = bool(visible)
        if self.search_panel_visible:
            if not self.search_panel.winfo_manager():
                self.search_panel.pack(fill="x", pady=(0, 8), before=self.text_frame)
        else:
            self.search_panel.pack_forget()

    def hide_search_panel(self):
        self._set_search_panel_visible(False)
        self._clear_search_highlights()
        self.text.focus_set()

    def open_search_panel(self, replace_mode=False):
        self._set_search_panel_visible(True)
        selected = self._selected_text_token()
        if selected and "\n" not in selected:
            self.search_var.set(selected)
        self._refresh_search_matches()
        if replace_mode:
            self.entry_replace.focus_set()
            self.entry_replace.selection_range(0, "end")
        else:
            self.entry_search.focus_set()
            self.entry_search.selection_range(0, "end")

    def _on_find_shortcut(self, _event):
        if _event is not None and (_event.state & 0x0001):
            return None
        self.open_search_panel(replace_mode=False)
        return "break"

    def _on_replace_shortcut(self, _event):
        if _event is not None and (_event.state & 0x0001):
            return None
        self.open_search_panel(replace_mode=True)
        return "break"

    def _on_search_next_shortcut(self, _event):
        self.search_next()
        return "break"

    def _on_search_prev_shortcut(self, _event):
        self.search_prev()
        return "break"

    def _on_replace_one_shortcut(self, _event):
        self.replace_current_match()
        return "break"

    def _clear_search_highlights(self):
        self.text.tag_remove(SEARCH_MATCH_TAG, "1.0", "end")
        self.text.tag_remove(SEARCH_CURRENT_TAG, "1.0", "end")
        self._search_match_ranges = []
        self._search_current_idx = -1

    def _find_search_matches(self):
        needle = self.search_var.get()
        if not needle:
            return []

        matches = []
        start = "1.0"
        while True:
            pos = self.text.search(
                needle,
                start,
                stopindex="end",
                exact=True,
                nocase=0 if self.search_case_var.get() else 1,
            )
            if not pos:
                break
            end = self.text.index(f"{pos}+{len(needle)}c")
            matches.append((pos, end))
            start = end
        return matches

    def _focus_current_search_match(self):
        self.text.tag_remove(SEARCH_CURRENT_TAG, "1.0", "end")
        if not self._search_match_ranges:
            return
        if self._search_current_idx < 0 or self._search_current_idx >= len(
            self._search_match_ranges
        ):
            self._search_current_idx = 0

        start, end = self._search_match_ranges[self._search_current_idx]
        self.text.tag_add(SEARCH_CURRENT_TAG, start, end)
        self.text.tag_raise(SEARCH_CURRENT_TAG)
        self.text.tag_raise(IDENTICAL_MATCH_TAG)
        self.text.mark_set("insert", end)
        self.text.see(start)
        self.text.tag_remove("sel", "1.0", "end")
        self.text.tag_add("sel", start, end)

    def _refresh_search_matches(self):
        needle = self.search_var.get()
        self.text.tag_remove(SEARCH_MATCH_TAG, "1.0", "end")
        self.text.tag_remove(SEARCH_CURRENT_TAG, "1.0", "end")
        self._search_match_ranges = []
        self._search_current_idx = -1

        if not needle:
            self._set_status("Søkefelt tomt.")
            return

        self._search_match_ranges = self._find_search_matches()
        for start, end in self._search_match_ranges:
            self.text.tag_add(SEARCH_MATCH_TAG, start, end)
        self.text.tag_raise(SEARCH_MATCH_TAG)
        self.text.tag_raise(IDENTICAL_MATCH_TAG)

        if self._search_match_ranges:
            insert_idx = self.text.index("insert")
            self._search_current_idx = 0
            for i, (start, _end) in enumerate(self._search_match_ranges):
                if self.text.compare(start, ">=", insert_idx):
                    self._search_current_idx = i
                    break
            self._focus_current_search_match()
            self._set_status(f"Fant {len(self._search_match_ranges)} treff.")
        else:
            self._set_status("Fant ingen treff.")

    def _on_search_query_changed(self, _event):
        self._refresh_search_matches()

    def search_next(self):
        if not self.search_var.get():
            self.open_search_panel(replace_mode=False)
            return

        if not self._search_match_ranges:
            self._refresh_search_matches()
        if not self._search_match_ranges:
            return

        self._search_current_idx = (self._search_current_idx + 1) % len(
            self._search_match_ranges
        )
        self._focus_current_search_match()

    def search_prev(self):
        if not self.search_var.get():
            self.open_search_panel(replace_mode=False)
            return

        if not self._search_match_ranges:
            self._refresh_search_matches()
        if not self._search_match_ranges:
            return

        self._search_current_idx = (self._search_current_idx - 1) % len(
            self._search_match_ranges
        )
        self._focus_current_search_match()

    def replace_current_match(self):
        replacement = self.replace_var.get()
        if not self.search_var.get():
            self.open_search_panel(replace_mode=True)
            return

        if not self._search_match_ranges:
            self._refresh_search_matches()
        if not self._search_match_ranges:
            return

        start, end = self._search_match_ranges[self._search_current_idx]
        self.text.delete(start, end)
        if replacement:
            self.text.insert(start, replacement)
            self.text.mark_set("insert", f"{start}+{len(replacement)}c")
        else:
            self.text.mark_set("insert", start)

        self._dirty = True
        self._schedule_spellcheck()
        self._refresh_search_matches()

    def replace_all_matches(self):
        replacement = self.replace_var.get()
        if not self.search_var.get():
            self.open_search_panel(replace_mode=True)
            return

        if not self._search_match_ranges:
            self._refresh_search_matches()
        if not self._search_match_ranges:
            return

        count = len(self._search_match_ranges)
        for start, end in reversed(self._search_match_ranges):
            self.text.delete(start, end)
            if replacement:
                self.text.insert(start, replacement)

        self._dirty = True
        self._schedule_spellcheck()
        self._refresh_search_matches()
        self._set_status(f"Ersattet {count} treff.")

    def _line_is_heading_candidate(self, line_no):
        line_start = f"{line_no}.0"
        line_end = f"{line_no}.end"
        line_text = self.text.get(line_start, line_end).strip()
        if not line_text:
            return None
        if self._parse_bullet_line(line_text):
            return None
        if TOC_HEADING_PATTERN.match(line_text):
            return line_text
        if len(line_text) > 110:
            return None

        total_letters = 0
        bold_letters = 0
        raw_line = self.text.get(line_start, line_end)
        for i, ch in enumerate(raw_line):
            if not ch.isalpha():
                continue
            total_letters += 1
            idx = self.text.index(f"{line_start}+{i}c")
            if BOLD_TAG in self.text.tag_names(idx):
                bold_letters += 1

        if total_letters >= 3 and bold_letters / max(total_letters, 1) >= 0.8:
            return line_text
        return None

    def _build_toc_entries(self):
        entries = []
        line_count = int(self.text.index("end-1c").split(".")[0])
        for line_no in range(1, line_count + 1):
            heading = self._line_is_heading_candidate(line_no)
            if heading:
                entries.append((line_no, heading))
        return entries

    def _refresh_toc_listbox(self):
        if self._toc_listbox is None or not self._toc_listbox.winfo_exists():
            return
        self._toc_entries = self._build_toc_entries()
        self._toc_listbox.delete(0, "end")
        for line_no, heading in self._toc_entries:
            compact = heading if len(heading) <= 72 else heading[:69] + "..."
            self._toc_listbox.insert("end", f"{line_no:>3}  {compact}")

    def _jump_to_selected_toc_entry(self, _event=None):
        if self._toc_listbox is None or not self._toc_listbox.winfo_exists():
            return
        selection = self._toc_listbox.curselection()
        if not selection:
            return
        idx = int(selection[0])
        if idx < 0 or idx >= len(self._toc_entries):
            return
        line_no, heading = self._toc_entries[idx]
        target = f"{line_no}.0"
        self.text.mark_set("insert", target)
        self.text.see(target)
        self.text.focus_set()
        self._set_status(f"Hoppet til linje {line_no}: {heading[:40]}")

    def show_toc(self):
        if self._toc_window is not None and self._toc_window.winfo_exists():
            self._refresh_toc_listbox()
            self._toc_window.deiconify()
            self._toc_window.lift()
            self._toc_window.focus_force()
            return

        self._toc_window = tk.Toplevel(self.root)
        self._toc_window.title("Innhold")
        self._toc_window.geometry("420x420+720+60")
        self._toc_window.attributes("-topmost", True)
        self._toc_window.configure(bg=self.theme["bg"])

        panel = tk.Frame(
            self._toc_window,
            bg=self.theme["card"],
            highlightthickness=1,
            highlightbackground=self.theme["border"],
            highlightcolor=self.theme["border"],
        )
        panel.pack(fill="both", expand=True, padx=10, pady=10)

        top = tk.Frame(panel, bg=self.theme["card"])
        top.pack(fill="x")
        tk.Label(
            top,
            text="Innholdsfortegnelse",
            bg=self.theme["card"],
            fg=self.theme["text"],
            font=self.title_font,
            padx=8,
            pady=6,
        ).pack(side="left")

        btn_refresh = self._create_button(top, "Oppdater", self._refresh_toc_listbox)
        btn_refresh.pack(side="right", padx=6)

        list_frame = tk.Frame(panel, bg=self.theme["card"])
        list_frame.pack(fill="both", expand=True, padx=6, pady=(0, 6))

        self._toc_listbox = tk.Listbox(
            list_frame,
            font=self.ui_font,
            bg=self.theme["card"],
            fg=self.theme["text"],
            selectbackground="#d8e6ff",
            relief="flat",
            borderwidth=0,
            activestyle="none",
        )
        self._toc_listbox.pack(side="left", fill="both", expand=True)
        self._toc_listbox.bind("<Double-Button-1>", self._jump_to_selected_toc_entry)
        self._toc_listbox.bind("<Return>", self._jump_to_selected_toc_entry)

        toc_scroll = tk.Scrollbar(list_frame, command=self._toc_listbox.yview)
        toc_scroll.pack(side="right", fill="y")
        self._toc_listbox.configure(yscrollcommand=toc_scroll.set)

        btn_jump = self._create_button(panel, "Gå til", self._jump_to_selected_toc_entry)
        btn_jump.pack(anchor="e", padx=6, pady=(0, 8))

        self._refresh_toc_listbox()

    def _on_toc_shortcut(self, _event):
        self.show_toc()
        return "break"

    def _on_ctrl_backspace(self, _event):
        try:
            selection_start = self.text.index("sel.first")
            selection_end = self.text.index("sel.last")
            self.text.delete(selection_start, selection_end)
            return "break"
        except tk.TclError:
            pass

        before_cursor = self.text.get("1.0", "insert")
        if not before_cursor:
            return "break"

        cut_index = len(before_cursor)
        while cut_index > 0 and before_cursor[cut_index - 1].isspace():
            cut_index -= 1
        while cut_index > 0 and not before_cursor[cut_index - 1].isspace():
            cut_index -= 1

        delete_count = len(before_cursor) - cut_index
        if delete_count > 0:
            self.text.delete(f"insert-{delete_count}c", "insert")

        return "break"

    def _on_bold_shortcut(self, _event):
        self.toggle_bold()
        return "break"

    def _refresh_bold_button(self):
        label = "B*" if self.bold_typing_mode else "B"
        self.btn_bold.configure(text=label)

    def _insert_plain_with_bold(self, text, insert_at):
        if not text:
            return insert_at
        self.text.insert(insert_at, text)
        run_end = self.text.index(f"{insert_at}+{len(text)}c")
        self.text.tag_add(BOLD_TAG, insert_at, run_end)
        self.text.mark_set("insert", run_end)
        return run_end

    def _is_navigation_or_modifier_key(self, keysym):
        return keysym in {
            "BackSpace",
            "Delete",
            "Left",
            "Right",
            "Up",
            "Down",
            "Home",
            "End",
            "Page_Up",
            "Page_Down",
            "Insert",
            "Escape",
            "Tab",
            "Return",
            "KP_Enter",
            "Shift_L",
            "Shift_R",
            "Control_L",
            "Control_R",
            "Alt_L",
            "Alt_R",
            "Meta_L",
            "Meta_R",
            "Super_L",
            "Super_R",
            "Caps_Lock",
            "Num_Lock",
            "Scroll_Lock",
        }

    def _on_text_keypress(self, event):
        if not self.bold_typing_mode:
            return None

        # Ignore keyboard shortcuts, but still allow AltGr (Ctrl+Alt) characters.
        ctrl_down = bool(event.state & 0x0004)
        alt_down = bool(event.state & 0x0008)
        if ctrl_down and not alt_down:
            return None

        if self._is_navigation_or_modifier_key(event.keysym):
            return None

        if not event.char or ord(event.char) < 32:
            return None

        try:
            selection_start = self.text.index("sel.first")
            selection_end = self.text.index("sel.last")
            self.text.delete(selection_start, selection_end)
            insert_at = selection_start
        except tk.TclError:
            insert_at = self.text.index("insert")

        self._insert_plain_with_bold(event.char, insert_at)
        self.text.see("insert")
        self._dirty = True
        self._schedule_spellcheck()
        return "break"

    def _is_superscript_candidate_char(self, ch):
        if not ch or len(ch) != 1:
            return False
        if ch.isspace():
            return False
        if ch in "{}[]()":
            return False
        return ch.isalnum() or ch in "+-=*/.,"

    def _auto_format_superscript_at_cursor(self):
        insert_at = self.text.index("insert")

        if self.text.compare(insert_at, "<=", "1.0"):
            return False

        close_char_start = self.text.index(f"{insert_at}-1c")
        if self.text.get(close_char_start, insert_at) == "}":
            line_start = self.text.index("insert linestart")
            marker = self.text.search(
                "^{", close_char_start, stopindex=line_start, backwards=True
            )
            if marker:
                content_start = self.text.index(f"{marker}+2c")
                if self.text.compare(content_start, "<", close_char_start):
                    content = self.text.get(content_start, close_char_start)
                    if content and "\n" not in content:
                        self.text.delete(close_char_start, insert_at)
                        self.text.delete(marker, f"{marker}+2c")
                        sup_end = self.text.index(f"{marker}+{len(content)}c")
                        self.text.tag_add(SUPERSCRIPT_TAG, marker, sup_end)
                        self.text.mark_set("insert", sup_end)
                        return True

        if self.text.compare(insert_at, "<=", "1.1"):
            return False

        marker_start = self.text.index(f"{insert_at}-2c")
        marker = self.text.get(marker_start, f"{marker_start}+1c")
        candidate = self.text.get(f"{marker_start}+1c", insert_at)
        if marker != "^":
            return False
        if not self._is_superscript_candidate_char(candidate):
            return False

        self.text.delete(marker_start, f"{marker_start}+1c")
        sup_end = self.text.index(f"{marker_start}+1c")
        self.text.tag_add(SUPERSCRIPT_TAG, marker_start, sup_end)
        self.text.mark_set("insert", sup_end)
        return True

    def _on_text_keyrelease(self, event):
        # Ignore keyboard shortcuts, but still allow AltGr-produced characters.
        ctrl_down = bool(event.state & 0x0004)
        alt_down = bool(event.state & 0x0008)
        if ctrl_down and not alt_down:
            return None

        if self._is_navigation_or_modifier_key(event.keysym):
            return None

        if self._auto_format_superscript_at_cursor():
            self._dirty = True
            self._schedule_spellcheck()

        return None

    def _selection_is_fully_bold(self, start, end):
        index = self.text.index(start)
        while self.text.compare(index, "<", end):
            if BOLD_TAG not in self.text.tag_names(index):
                return False
            index = self.text.index(f"{index}+1c")
        return True

    def toggle_bold(self):
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
        except tk.TclError:
            self.bold_typing_mode = not self.bold_typing_mode
            self._refresh_bold_button()
            if self.bold_typing_mode:
                self._set_status("Bold skrivemodus PÅ. Trykk Ctrl+B for å slå av.")
            else:
                self._set_status("Bold skrivemodus AV.")
            return

        if self._selection_is_fully_bold(start, end):
            self.text.tag_remove(BOLD_TAG, start, end)
            self._set_status("Bold AV.")
        else:
            self.text.tag_add(BOLD_TAG, start, end)
            self._set_status("Bold PÅ.")

        self._dirty = True
        self.text.see("insert")

    def _on_bullet_shortcut(self, _event):
        self.toggle_bullets()
        return "break"

    def _on_mark_identical_shortcut(self, _event):
        self.mark_identical()
        return "break"

    def _on_replace_identical_shortcut(self, _event):
        self.replace_identical()
        return "break"

    def _on_spellcheck_shortcut(self, _event):
        self.toggle_spellcheck()
        return "break"

    def _clear_identical_highlight(self):
        self.text.tag_remove(IDENTICAL_MATCH_TAG, "1.0", "end")

    def _selected_text_token(self):
        try:
            return self.text.get("sel.first", "sel.last")
        except tk.TclError:
            return ""

    def _short_token_label(self, token, max_len=18):
        compact = token.replace("\r", " ").replace("\n", " ")
        if len(compact) <= max_len:
            return compact
        return compact[: max_len - 3] + "..."

    def _find_exact_occurrences(self, token):
        if not token:
            return []

        matches = []
        token_len = len(token)
        search_from = "1.0"

        while True:
            start = self.text.search(token, search_from, stopindex="end", exact=True)
            if not start:
                break

            end = self.text.index(f"{start}+{token_len}c")
            matches.append((start, end))
            search_from = end

        return matches

    def _highlight_exact_occurrences(self, token):
        self._clear_identical_highlight()
        matches = self._find_exact_occurrences(token)
        for start, end in matches:
            self.text.tag_add(IDENTICAL_MATCH_TAG, start, end)
        self.text.tag_raise(IDENTICAL_MATCH_TAG)
        return matches

    def mark_identical(self):
        token = self._selected_text_token()
        if not token:
            self._set_status("Marker tegn/tekst du vil finne like forekomster av.")
            return

        self._last_identical_token = token
        matches = self._highlight_exact_occurrences(token)
        label = self._short_token_label(token)
        self._set_status(
            f"Markerte {len(matches)} like treff for '{label}'. Erstatt: Ctrl+Shift+R."
        )

    def replace_identical(self):
        token = self._selected_text_token()
        if token:
            self._last_identical_token = token
        else:
            token = self._last_identical_token

        if not token:
            self._set_status("Marker tegn/tekst du vil erstatte eller slette.")
            return

        matches = self._find_exact_occurrences(token)
        if not matches:
            self._set_status("Fant ingen like treff å endre.")
            return

        label = self._short_token_label(token)
        replacement = simpledialog.askstring(
            "Erstatt like",
            f"Ny tekst for '{label}'.\nLa feltet stå tomt for å slette alle treff.",
            initialvalue=token,
            parent=self.root,
        )
        if replacement is None:
            self._set_status("Erstatting avbrutt.")
            return

        if replacement == token:
            highlighted = self._highlight_exact_occurrences(token)
            self._set_status(f"Ingen endring. Markerte {len(highlighted)} treff.")
            return

        for start, end in reversed(matches):
            self.text.delete(start, end)
            if replacement:
                self.text.insert(start, replacement)

        self._dirty = True
        self._schedule_spellcheck()

        if replacement:
            self._last_identical_token = replacement
            highlighted = self._highlight_exact_occurrences(replacement)
            self._set_status(
                f"Ersattet {len(matches)} treff. Markerte {len(highlighted)} nye treff."
            )
        else:
            self._last_identical_token = ""
            self._clear_identical_highlight()
            self._set_status(f"Slettet {len(matches)} like treff.")

    def _bullet_level_from_indent(self, indent_text):
        width = 0
        for ch in indent_text:
            if ch == "\t":
                width += BULLET_INDENT_SPACES - (width % BULLET_INDENT_SPACES)
            else:
                width += 1
        return max(0, width // BULLET_INDENT_SPACES)

    def _bullet_marker_for_level(self, level):
        if level < 0:
            level = 0
        return BULLET_MARKERS[min(level, len(BULLET_MARKERS) - 1)]

    def _bullet_prefix_for_level(self, level):
        return (" " * (max(0, level) * BULLET_INDENT_SPACES)) + self._bullet_marker_for_level(
            level
        ) + " "

    def _parse_bullet_line(self, line_text):
        match = BULLET_LINE_PATTERN.match(line_text)
        if not match:
            return None

        indent = match.group(1)
        marker = match.group(2)
        prefix_len = match.end()
        level = self._bullet_level_from_indent(indent)
        return {
            "indent": indent,
            "marker": marker,
            "prefix_len": prefix_len,
            "level": level,
            "content": line_text[prefix_len:],
        }

    def _selected_line_range(self):
        try:
            sel_start = self.text.index("sel.first")
            sel_end = self.text.index("sel.last")
            start_line = int(sel_start.split(".")[0])
            end_line = int(sel_end.split(".")[0])
            end_col = int(sel_end.split(".")[1])
            if end_col == 0 and end_line > start_line:
                end_line -= 1
        except tk.TclError:
            start_line = int(self.text.index("insert").split(".")[0])
            end_line = start_line

        if start_line > end_line:
            start_line, end_line = end_line, start_line
        return start_line, end_line

    def _on_return(self, _event):
        line_start = self.text.index("insert linestart")
        line_end = self.text.index("insert lineend")
        line_text = self.text.get(line_start, line_end)
        bullet_info = self._parse_bullet_line(line_text)
        if not bullet_info:
            return None

        content = bullet_info["content"]
        if content.strip() == "":
            marker_start = f"{line_start}+{len(bullet_info['indent'])}c"
            marker_end = f"{line_start}+{bullet_info['prefix_len']}c"
            self.text.delete(marker_start, marker_end)
            return None

        insert_at = self.text.index("insert")
        if self.text.compare(insert_at, "==", line_end):
            trimmed_content = content.rstrip()
            if trimmed_content and trimmed_content[-1] not in ".!?:":
                trailing_spaces = len(content) - len(trimmed_content)
                if trailing_spaces > 0:
                    dot_index = f"{line_end}-{trailing_spaces}c"
                else:
                    dot_index = line_end
                self.text.insert(dot_index, ".")
            self.text.mark_set("insert", "insert lineend")

        next_prefix = bullet_info["indent"] + bullet_info["marker"] + " "
        self.text.insert("insert", "\n" + next_prefix)
        self._dirty = True
        self._schedule_spellcheck()
        return "break"

    def _toggle_bullet_line(self, line_no):
        line_start = f"{line_no}.0"
        line_end = f"{line_no}.end"
        line_text = self.text.get(line_start, line_end)
        bullet_info = self._parse_bullet_line(line_text)

        if bullet_info:
            marker_start = f"{line_start}+{len(bullet_info['indent'])}c"
            marker_end = f"{line_start}+{bullet_info['prefix_len']}c"
            self.text.delete(marker_start, marker_end)
            return

        indent = re.match(r"^[ \t]*", line_text).group(0)
        level = self._bullet_level_from_indent(indent)
        marker = self._bullet_marker_for_level(level)
        insert_at = f"{line_start}+{len(indent)}c"
        self.text.insert(insert_at, marker + " ")

    def _change_bullet_level(self, delta):
        start_line, end_line = self._selected_line_range()
        changed = 0

        for line_no in range(start_line, end_line + 1):
            line_start = f"{line_no}.0"
            line_end = f"{line_no}.end"
            line_text = self.text.get(line_start, line_end)
            bullet_info = self._parse_bullet_line(line_text)
            if not bullet_info:
                continue

            new_level = max(0, bullet_info["level"] + delta)
            if new_level == bullet_info["level"]:
                continue

            new_prefix = self._bullet_prefix_for_level(new_level)
            old_prefix_end = f"{line_start}+{bullet_info['prefix_len']}c"
            self.text.delete(line_start, old_prefix_end)
            self.text.insert(line_start, new_prefix)
            changed += 1

        if changed > 0:
            self._dirty = True
            self._schedule_spellcheck()
            self.text.see("insert")

        return changed

    def _on_tab_indent(self, _event):
        changed = self._change_bullet_level(+1)
        if changed > 0:
            self._set_status(f"Indent økte for {changed} punkt.")
            return "break"
        return None

    def _on_shift_tab_outdent(self, _event):
        changed = self._change_bullet_level(-1)
        if changed > 0:
            self._set_status(f"Indent senket for {changed} punkt.")
            return "break"
        return None

    def toggle_bullets(self):
        start_line, end_line = self._selected_line_range()
        for line_no in range(start_line, end_line + 1):
            self._toggle_bullet_line(line_no)

        self._dirty = True
        self.text.see("insert")
        self._set_status(
            "Bullet oppdatert. Shortcut: Ctrl+Shift+B. Nivå: Tab / Shift+Tab."
        )

    def _refresh_spellcheck_button(self):
        if not self._spellcheck_available:
            self.spellcheck_enabled = False
            self.btn_spellcheck.configure(text="Stv: mangler", state="disabled")
            return

        label = "Stv: PÅ" if self.spellcheck_enabled else "Stv: AV"
        self.btn_spellcheck.configure(text=label, state="normal")

    def toggle_spellcheck(self):
        if not self._spellcheck_available:
            self._set_status("Stavekontroll krever pakken wordfreq.")
            return

        self.spellcheck_enabled = not self.spellcheck_enabled
        self._refresh_spellcheck_button()

        if self.spellcheck_enabled:
            self._set_status("Stavekontroll PÅ.")
            self._schedule_spellcheck()
            return

        if self._spellcheck_after_id is not None:
            self.root.after_cancel(self._spellcheck_after_id)
            self._spellcheck_after_id = None
        self.text.tag_remove(SPELLCHECK_TAG, "1.0", "end")
        self._set_status("Stavekontroll AV.")

    def _word_looks_correct(self, word):
        if not self._spellcheck_available:
            return True

        if len(word) <= 1:
            return True

        lowered = word.casefold()
        if lowered in self.custom_spell_words:
            return True

        if any(ch.isdigit() for ch in lowered):
            return True

        if word.isupper() and len(word) <= 6:
            return True

        if lowered.startswith("http") or lowered.startswith("www"):
            return True

        if self._looks_like_hyphenated_name(word):
            return True

        best_score = self._best_zipf_score(lowered)
        if best_score is None:
            return True

        required_score = self._required_zipf_score(lowered)
        if best_score < required_score:
            if self._looks_like_compound_word(lowered):
                return True
            return False

        # Common informal shorting: "happenin" -> "happening".
        if lowered.endswith("in") and len(lowered) >= 5 and lowered.isalpha():
            ing_scores = []
            ing_variant = lowered + "g"
            for lang in SPELLCHECK_LANGS:
                try:
                    ing_scores.append(float(zipf_frequency(ing_variant, lang)))
                except Exception:
                    continue

            if ing_scores and max(ing_scores) >= max(required_score, best_score + 1.0):
                return False

        return True

    def _looks_like_hyphenated_name(self, word):
        if "-" not in word:
            return False

        parts = [part for part in re.split(r"[-\u2010-\u2015]", word) if part]
        if len(parts) < 2:
            return False

        if not all(part.isalpha() and len(part) >= 2 for part in parts):
            return False

        for part in parts:
            if part.isupper():
                continue
            if not part[0].isupper():
                return False

        strong_parts = 0
        for part in parts:
            score = self._best_zipf_score(part.casefold()) or 0.0
            if score >= 1.3:
                strong_parts += 1
        return strong_parts >= 1

    def _best_zipf_score(self, lowered_word):
        best_score = None
        for lang in SPELLCHECK_LANGS:
            try:
                score = float(zipf_frequency(lowered_word, lang))
            except Exception:
                continue
            if best_score is None or score > best_score:
                best_score = score
        return best_score

    def _compound_segment_score(self, segment, side):
        variants = {segment}
        if len(segment) < 3:
            return 0.0

        # Compound glue letter "s" (e.g. sesjon+s+nokkel).
        if side == "left" and segment.endswith("s") and len(segment) > 4:
            variants.add(segment[:-1])
        if side == "right" and segment.startswith("s") and len(segment) > 4:
            variants.add(segment[1:])

        # Common Norwegian inflection endings. Useful when wordfreq has
        # one form but not another.
        for suffix in ("ene", "ane", "ens", "ets", "ers", "er", "en", "et", "ar", "e", "a"):
            if segment.endswith(suffix) and len(segment) - len(suffix) >= 3:
                variants.add(segment[: -len(suffix)])

        best = 0.0
        for candidate in variants:
            score = self._best_zipf_score(candidate)
            if score is not None and score > best:
                best = score
        return best

    def _looks_like_compound_word(self, lowered_word):
        if len(lowered_word) < 8 or not lowered_word.isalpha():
            return False

        # Norwegian compound words are often missing as full-form entries in wordfreq.
        # Accept if the word can be split into two reasonably common chunks.
        for split_at in range(3, len(lowered_word) - 2):
            left = lowered_word[:split_at]
            right = lowered_word[split_at:]
            left_score = self._compound_segment_score(left, "left")
            right_score = self._compound_segment_score(right, "right")

            if left_score >= 2.2 and right_score >= 2.2:
                return True

            # Allow one rarer chunk when the other part is clearly common.
            if len(lowered_word) >= 11 and left_score >= 2.8 and right_score >= 1.6:
                return True
            if len(lowered_word) >= 11 and right_score >= 2.8 and left_score >= 1.6:
                return True
        return False

    def _required_zipf_score(self, lowered_word):
        length = len(lowered_word)
        if length <= 2:
            return 5.0
        if length <= 3:
            return 4.3
        if length <= 4:
            return 3.7
        if length <= 6:
            return 2.8
        if length <= 8:
            return 2.1
        if length <= 10:
            return 1.9
        if length <= 14:
            return 1.6
        return 1.45

    def _schedule_spellcheck(self):
        if self._spellcheck_after_id is not None:
            self.root.after_cancel(self._spellcheck_after_id)
            self._spellcheck_after_id = None

        if not self.spellcheck_enabled or not self._spellcheck_available:
            return

        self._spellcheck_after_id = self.root.after(
            SPELLCHECK_DEBOUNCE_MS, self._run_spellcheck
        )

    def _run_spellcheck(self):
        self._spellcheck_after_id = None
        self.text.tag_remove(SPELLCHECK_TAG, "1.0", "end")

        if not self.spellcheck_enabled or not self._spellcheck_available:
            return

        line_count = int(self.text.index("end-1c").split(".")[0])
        for line_no in range(1, line_count + 1):
            line_start = f"{line_no}.0"
            line_end = f"{line_no}.end"
            line_text = self.text.get(line_start, line_end)

            for match in SPELLCHECK_WORD_PATTERN.finditer(line_text):
                start = f"{line_no}.{match.start()}"
                end = f"{line_no}.{match.end()}"
                if IMAGE_TOKEN_HIDDEN_TAG in self.text.tag_names(start):
                    continue
                if SUPERSCRIPT_TAG in self.text.tag_names(start):
                    continue
                if self._word_looks_correct(match.group(0)):
                    continue
                self.text.tag_add(SPELLCHECK_TAG, start, end)

    def _set_status(self, message):
        self.status_var.set(message)

    def _hide_window(self):
        self.root.withdraw()
        self.root.attributes("-alpha", self._target_alpha())

    def _refresh_auto_capture_button(self):
        label = "Auto: PÅ" if self.auto_capture_enabled else "Auto: AV"
        self.btn_auto_capture.configure(text=label)

    def toggle_auto_capture(self):
        self.auto_capture_enabled = not self.auto_capture_enabled
        if self.auto_capture_enabled:
            current_capture = self._read_clipboard_capture()
            self._last_clipboard_signature = (
                current_capture["signature"] if current_capture else None
            )
            self._set_status("Auto-fangst PÅ. Marker tekst i andre vinduer og trykk Ctrl+C.")
        else:
            self._set_status("Auto-fangst AV.")
        self._refresh_auto_capture_button()

    def _read_clipboard_text(self):
        try:
            text = self.root.clipboard_get()
        except tk.TclError:
            return None

        if not isinstance(text, str):
            return None

        cleaned = self._normalize_captured_text(text).strip()
        return cleaned if cleaned else None

    def _clipboard_format_name(self, format_id):
        standard_name = CLIPBOARD_STANDARD_FORMATS.get(format_id)
        if standard_name:
            return standard_name

        user32 = self._u32
        buffer = ctypes.create_unicode_buffer(256)
        count = user32.GetClipboardFormatNameW(format_id, buffer, 256)
        if count > 0:
            return buffer.value
        return f"FORMAT_{format_id}"

    def _clipboard_available_format_names(self):
        user32 = self._u32
        if not self._open_clipboard_with_retry():
            return []

        names = []
        try:
            fmt = 0
            while True:
                fmt = user32.EnumClipboardFormats(fmt)
                if fmt == 0:
                    break
                names.append(self._clipboard_format_name(fmt))
        finally:
            self._close_clipboard_quietly()

        return names


    def _read_clipboard_format_bytes(self, format_names, keyword=None):
        user32 = self._u32
        kernel32 = self._k32
        if isinstance(format_names, str):
            format_names = [format_names]

        format_ids = []
        for format_name in format_names:
            format_id = user32.RegisterClipboardFormatW(format_name)
            if format_id:
                format_ids.append(format_id)

        if not format_ids and not keyword:
            return None

        if not self._open_clipboard_with_retry():
            return None


        raw_data = None
        try:
            selected_format_id = None
            for format_id in format_ids:
                if user32.IsClipboardFormatAvailable(format_id):
                    selected_format_id = format_id
                    break

            if selected_format_id is None and keyword:
                current = 0
                while True:
                    current = user32.EnumClipboardFormats(current)
                    if current == 0:
                        break
                    format_name = self._clipboard_format_name(current).casefold()
                    if keyword.casefold() in format_name:
                        selected_format_id = current
                        break

            if selected_format_id is None:
                return None

            handle = user32.GetClipboardData(selected_format_id)
            if not handle:
                return None

            size = int(kernel32.GlobalSize(handle))
            if size <= 0:
                return None

            ptr = kernel32.GlobalLock(handle)
            if not ptr:
                return None

            try:
                raw_data = ctypes.string_at(ptr, size)
            finally:
                kernel32.GlobalUnlock(handle)
        finally:
            self._close_clipboard_quietly()


        if not raw_data:
            return None
        return raw_data

    def _decode_clipboard_bytes(self, raw_data):
        if not raw_data:
            return ""

        data = raw_data

        def looks_like_utf16le(blob):
            sample = blob[: min(len(blob), 1024)]
            if len(sample) < 8:
                return False
            even = sample[0::2]
            odd = sample[1::2]
            if not even or not odd:
                return False
            even_zero = even.count(0) / len(even)
            odd_zero = odd.count(0) / len(odd)
            return odd_zero > 0.55 and even_zero < 0.25

        def looks_like_utf16be(blob):
            sample = blob[: min(len(blob), 1024)]
            if len(sample) < 8:
                return False
            even = sample[0::2]
            odd = sample[1::2]
            if not even or not odd:
                return False
            even_zero = even.count(0) / len(even)
            odd_zero = odd.count(0) / len(odd)
            return even_zero > 0.55 and odd_zero < 0.25

        looks_le = looks_like_utf16le(data)
        looks_be = looks_like_utf16be(data)
        starts_with_bom = data.startswith(b"\xff\xfe") or data.startswith(b"\xfe\xff")

        # Trim padding nulls often present in HGLOBAL buffers.
        if not (starts_with_bom or looks_le or looks_be):
            data = re.sub(rb"\x00+$", b"", data)

        if starts_with_bom:
            encodings = ("utf-16", "utf-8", "cp1252", "latin1")
        elif looks_le:
            if len(data) % 2 == 1:
                data = data[:-1]
            encodings = ("utf-16le", "utf-16", "utf-8", "cp1252", "latin1")
        elif looks_be:
            if len(data) % 2 == 1:
                data = data[:-1]
            encodings = ("utf-16be", "utf-16", "utf-8", "cp1252", "latin1")
        else:
            encodings = ("utf-8", "cp1252", "latin1", "utf-16le")

        for encoding in encodings:
            try:
                decoded = data.decode(encoding)
            except Exception:
                continue

            if "\x00" in decoded:
                decoded = decoded.replace("\x00", "")
            return decoded

        decoded = data.decode("utf-8", errors="ignore")
        return decoded.replace("\x00", "")

    def _extract_html_fragment_from_text(self, html_text):
        if not html_text:
            return ""

        marker_match = re.search(
            r"<!--\s*StartFragment\s*-->(.*?)<!--\s*EndFragment\s*-->",
            html_text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if marker_match:
            return marker_match.group(1)

        offset_match = re.search(
            r"StartFragment:(\d+).*?EndFragment:(\d+)",
            html_text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if offset_match:
            start = int(offset_match.group(1))
            end = int(offset_match.group(2))
            if 0 <= start < end <= len(html_text):
                return html_text[start:end]

        body_match = re.search(
            r"<body[^>]*>(.*?)</body>", html_text, flags=re.IGNORECASE | re.DOTALL
        )
        if body_match:
            return body_match.group(1)

        return html_text

    def _read_clipboard_html_fragment(self):
        raw_data = self._read_clipboard_format_bytes(
            CLIPBOARD_HTML_FORMAT_NAMES, keyword="html"
        )
        if not raw_data:
            return None

        compact = re.sub(rb"\x00+$", b"", raw_data)
        html_bytes = compact
        fragment_bytes = compact

        start_html_match = re.search(rb"StartHTML:(\d+)", compact)
        end_html_match = re.search(rb"EndHTML:(\d+)", compact)
        if start_html_match and end_html_match:
            html_start = int(start_html_match.group(1))
            html_end = int(end_html_match.group(1))
            if 0 <= html_start < html_end <= len(compact):
                html_bytes = compact[html_start:html_end]

        start_match = re.search(rb"StartFragment:(\d+)", compact)
        end_match = re.search(rb"EndFragment:(\d+)", compact)
        if start_match and end_match:
            start = int(start_match.group(1))
            end = int(end_match.group(1))
            if 0 <= start < end <= len(compact):
                fragment_bytes = compact[start:end]
            else:
                fragment_bytes = html_bytes
        else:
            start_marker = b"<!--StartFragment-->"
            end_marker = b"<!--EndFragment-->"
            start = html_bytes.find(start_marker)
            end = html_bytes.find(end_marker)
            if start != -1 and end != -1 and end > start:
                fragment_bytes = html_bytes[start + len(start_marker) : end]
            else:
                fragment_bytes = html_bytes

        decoded_html = self._decode_clipboard_bytes(html_bytes)
        decoded_fragment = self._decode_clipboard_bytes(fragment_bytes)
        if decoded_fragment:
            style_maps = self._extract_html_style_maps(decoded_html)
            return {
                "fragment": decoded_fragment,
                "html_text": decoded_html,
                "class_bold_map": style_maps.get("class_bold_map", {}),
                "css_vars": style_maps.get("css_vars", {}),
            }

        decoded_html = self._decode_clipboard_bytes(raw_data)
        decoded_fragment = self._extract_html_fragment_from_text(decoded_html)
        if not decoded_fragment:
            decoded_fragment = decoded_html

        style_maps = self._extract_html_style_maps(decoded_html)
        return {
            "fragment": decoded_fragment,
            "html_text": decoded_html,
            "class_bold_map": style_maps.get("class_bold_map", {}),
            "css_vars": style_maps.get("css_vars", {}),
        }

    def _extract_html_style_maps(self, html_text):
        class_bold_map = {}
        css_vars = {}
        if not html_text:
            return {"class_bold_map": class_bold_map, "css_vars": css_vars}

        style_blocks = re.findall(
            r"<style[^>]*>(.*?)</style>", html_text, flags=re.IGNORECASE | re.DOTALL
        )

        for block in style_blocks:
            for _selector_text, declarations in re.findall(
                r"([^{}]+)\{([^{}]+)\}", block, flags=re.DOTALL
            ):
                for var_name, var_value in re.findall(
                    r"(--[A-Za-z0-9_-]+)\s*:\s*([^;]+)",
                    declarations,
                    flags=re.IGNORECASE,
                ):
                    css_vars[var_name.casefold()] = var_value.strip()

        for block in style_blocks:
            for selector_text, declarations in re.findall(
                r"([^{}]+)\{([^{}]+)\}", block, flags=re.DOTALL
            ):
                if not ClipboardHtmlRunParser._style_implies_bold(declarations, css_vars):
                    continue
                for class_name in re.findall(r"\.([A-Za-z0-9_-]+)", selector_text):
                    class_bold_map[class_name] = True

        return {"class_bold_map": class_bold_map, "css_vars": css_vars}

    def _read_clipboard_rtf_runs(self):
        raw_data = self._read_clipboard_format_bytes(
            CLIPBOARD_RTF_FORMAT_NAMES, keyword="rtf"
        )
        if not raw_data:
            return []

        rtf_text = self._decode_clipboard_bytes(raw_data)
        if not rtf_text:
            return []

        rtf_start = rtf_text.find(r"{\rtf")
        if rtf_start > 0:
            rtf_text = rtf_text[rtf_start:]

        parser = ClipboardRtfRunParser()
        try:
            runs = parser.parse(rtf_text)
        except Exception:
            return []

        return self._normalize_captured_runs(runs)

    def _normalize_captured_runs(self, runs):
        normalized = []
        for run_text, is_bold in runs:
            cleaned = self._normalize_captured_text(run_text)
            cleaned = cleaned.replace("\u2022", "-")
            cleaned = re.sub(r"\n[ \t]*-\s*\n", "\n- ", cleaned)
            cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
            cleaned = cleaned.replace("\n\n- ", "\n- ")
            cleaned = cleaned.replace("\n-  ", "\n- ")
            if not cleaned:
                continue

            if normalized and normalized[-1][1] == is_bold:
                normalized[-1][0] += cleaned
            else:
                normalized.append([cleaned, bool(is_bold)])

        if not normalized:
            return []

        for i in range(1, len(normalized)):
            prev = normalized[i - 1][0]
            curr = normalized[i][0]
            if prev.endswith(" ") and curr.startswith(" "):
                normalized[i][0] = curr.lstrip(" ")

        normalized = [item for item in normalized if item[0]]
        if not normalized:
            return []

        normalized[0][0] = normalized[0][0].lstrip()
        normalized[-1][0] = normalized[-1][0].rstrip()
        normalized = [item for item in normalized if item[0]]
        return [(text, is_bold) for text, is_bold in normalized]

    def _canonical_capture_text(self, text):
        if not text:
            return ""
        canonical = self._normalize_captured_text(text)
        canonical = canonical.replace("\u2022", "-")
        canonical = re.sub(r"\n[ \t]*-\s*", "\n", canonical)
        canonical = re.sub(r"\s+", " ", canonical)
        return canonical.strip().casefold()

    def _capture_similarity(self, text_a, text_b):
        if not text_a or not text_b:
            return 0.0
        return SequenceMatcher(None, text_a, text_b).ratio()

    def _read_clipboard_capture(self):
        plain_text = self._read_clipboard_text()
        canonical_plain_text = self._canonical_capture_text(plain_text)

        html_capture = None
        html_payload = self._read_clipboard_html_fragment()
        if html_payload:
            class_bold_map = html_payload.get("class_bold_map", {})
            css_vars = html_payload.get("css_vars", {})
            html_text = html_payload.get("html_text", "")

            fragment_candidates = []
            primary_fragment = html_payload.get("fragment", "")
            if primary_fragment:
                fragment_candidates.append(primary_fragment)

            alt_fragment = self._extract_html_fragment_from_text(html_text)
            if alt_fragment and alt_fragment not in fragment_candidates:
                fragment_candidates.append(alt_fragment)

            if html_text and html_text not in fragment_candidates:
                fragment_candidates.append(html_text)

            best_runs = None
            best_score = (-1, -1.0, -10**9, -1)
            best_similarity = 0.0

            for candidate in fragment_candidates:
                parser = ClipboardHtmlRunParser(class_bold_map, css_vars)
                try:
                    parser.feed(candidate)
                    parser.close()
                except Exception:
                    continue

                runs = self._normalize_captured_runs(parser.runs)
                if not runs:
                    continue

                candidate_text = "".join(text for text, _ in runs)
                if not candidate_text:
                    continue

                candidate_canonical = self._canonical_capture_text(candidate_text)
                similarity = self._capture_similarity(
                    candidate_canonical, canonical_plain_text
                )
                bold_chars = sum(len(text) for text, is_bold in runs if is_bold)
                length_delta = abs(len(candidate_canonical) - len(canonical_plain_text))
                score = (
                    1 if bold_chars > 0 else 0,
                    similarity,
                    -length_delta,
                    len(candidate_text),
                )
                if score > best_score:
                    best_score = score
                    best_similarity = similarity
                    best_runs = runs

            if best_runs and (not canonical_plain_text or best_similarity >= 0.28):
                html_capture = {
                    "text": "".join(text for text, _ in best_runs),
                    "runs": best_runs,
                    "signature": ("html", tuple(best_runs)),
                    "source": "html",
                }

        rtf_capture = None
        rtf_runs = self._read_clipboard_rtf_runs()
        if rtf_runs:
            rtf_text = "".join(text for text, _ in rtf_runs)
            if rtf_text:
                rtf_capture = {
                    "text": rtf_text,
                    "runs": rtf_runs,
                    "signature": ("rtf", tuple(rtf_runs)),
                    "source": "rtf",
                }

        if html_capture and any(is_bold for _, is_bold in html_capture["runs"]):
            return html_capture
        if rtf_capture and any(is_bold for _, is_bold in rtf_capture["runs"]):
            return rtf_capture
        if html_capture:
            return html_capture
        if rtf_capture:
            return rtf_capture

        text = plain_text
        if not text:
            return None

        available_formats = self._clipboard_available_format_names()
        lowered_formats = [name.casefold() for name in available_formats]
        has_html_like = any("html" in name for name in lowered_formats)
        has_rtf_like = any("rtf" in name for name in lowered_formats)
        visible_formats = ", ".join(available_formats[:6]) if available_formats else "none"
        if has_html_like or has_rtf_like:
            detail = f"fant html/rtf format men klarte ikke lese bold. formater: {visible_formats}"
        else:
            detail = f"ingen html/rtf format i clipboard. formater: {visible_formats}"

        return {
            "text": text,
            "runs": [(text, False)],
            "signature": ("text", text),
            "source": "text",
            "detail": detail,
        }

    def _normalize_captured_text(self, text):
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = normalized.replace("\u00A0", " ").replace("\u202F", " ")
        normalized_lines = [
            MULTISPACE_PATTERN.sub(" ", line) for line in normalized.split("\n")
        ]
        return "\n".join(normalized_lines)

    def _insert_runs_at(self, index, runs):
        cursor = index
        for run_text, is_bold in runs:
            if not run_text:
                continue
            self.text.insert(cursor, run_text)
            run_end = self.text.index(f"{cursor}+{len(run_text)}c")
            if is_bold:
                self.text.tag_add(BOLD_TAG, cursor, run_end)
            cursor = run_end
        return cursor

    def _append_captured_text(self, capture):
        try:
            selection_start = self.text.index("sel.first")
            selection_end = self.text.index("sel.last")
            self.text.delete(selection_start, selection_end)
            insert_at = selection_start
        except tk.TclError:
            insert_at = self.text.index("insert")

        runs = capture.get("runs") or [(capture.get("text", ""), False)]
        end_at = self._insert_runs_at(insert_at, runs)
        self.text.mark_set("insert", end_at)
        self.text.see("insert")
        self._dirty = True

    def _clipboard_watch_tick(self):
        if self.auto_capture_enabled:
            capture = self._read_clipboard_capture()
            if capture and capture["signature"] != self._last_clipboard_signature:
                self._last_clipboard_signature = capture["signature"]
                self._append_captured_text(capture)
                bold_runs = sum(1 for _text, is_bold in capture.get("runs", []) if is_bold)
                source = capture.get("source", "text")
                detail = capture.get("detail")
                if bold_runs > 0:
                    self._set_status(
                        f"Auto-fanget tekst fra clipboard ({source}, bold beholdt)."
                    )
                else:
                    if detail:
                        self._set_status(
                            f"Auto-fanget tekst fra clipboard ({source}: {detail})"
                        )
                    else:
                        self._set_status(f"Auto-fanget tekst fra clipboard ({source}).")

        self.root.after(CLIPBOARD_POLL_MS, self._clipboard_watch_tick)

    def _on_text_modified(self, _event):
        if self.text.edit_modified():
            self._dirty = True
            self._schedule_spellcheck()
            if self.search_panel_visible and self.search_var.get():
                self._refresh_search_matches()
            if self._toc_window is not None and self._toc_window.winfo_exists():
                self._refresh_toc_listbox()
            self.text.edit_modified(False)

    def _autosave_tick(self):
        if self._dirty:
            self._save_session(silent=True)
        self.root.after(self.autosave_ms, self._autosave_tick)

    def toggle_visibility(self):
        state = str(self.root.state())
        if state == "withdrawn":
            self.root.deiconify()
            self.root.geometry("+0+0")
            self.root.lift()
            self.root.focus_force()
            self.text.focus_set()
            self.root.attributes("-alpha", 0.0)
            self._fade_to(self._target_alpha())
            self._set_status("Notatvindu åpnet.")
        else:
            self._fade_to(0.0, on_complete=self._hide_window)

    def new_note(self):
        if not messagebox.askyesno("Nytt notat", "Tømme notatet og starte på nytt?"):
            return
        self.text.delete("1.0", "end")
        self._clear_identical_highlight()
        self._last_identical_token = ""
        self._clear_search_highlights()
        self.search_var.set("")
        self.replace_var.set("")
        self.attachments = {}
        self.formula_meta = {}
        self._inline_image_refs = []
        self.text.tag_remove(SPELLCHECK_TAG, "1.0", "end")
        self._dirty = True
        self._save_session(silent=True)
        self._set_status("Nytt notat opprettet.")

    def _on_paste(self, _event):
        image = self._clipboard_image()
        if image is None:
            return None
        self._save_image_and_insert_token(image)
        self._dirty = True
        self._set_status("Bilde limt inn.")
        return "break"

    def _clipboard_image(self):
        try:
            content = ImageGrab.grabclipboard()
        except Exception:
            return None

        if isinstance(content, Image.Image):
            return content

        if isinstance(content, list):
            for item in content:
                path = Path(item)
                if path.exists() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp"}:
                    try:
                        return Image.open(path)
                    except Exception:
                        continue
        return None

    def _normalize_latex_expression(self, raw_expression):
        text = (raw_expression or "").strip()
        if len(text) >= 4 and text.startswith("$$") and text.endswith("$$"):
            text = text[2:-2].strip()
        elif len(text) >= 2 and text.startswith("$") and text.endswith("$"):
            text = text[1:-1].strip()
        return text

    def _wrap_mathtext_expression(self, latex_expression):
        expression = self._normalize_latex_expression(latex_expression)
        if not expression:
            return ""
        if expression.startswith("$") and expression.endswith("$"):
            return expression
        return f"${expression}$"

    def _render_formula_to_image(self, latex_expression, target_path):
        if FigureCanvasAgg is None or Figure is None:
            raise RuntimeError("Matplotlib mangler. Kjør: pip install matplotlib")

        expression = self._wrap_mathtext_expression(latex_expression)
        if not expression:
            raise ValueError("Tom formel.")

        fig = Figure(figsize=(6.0, 1.6), dpi=FORMULA_RENDER_DPI)
        try:
            fig.patch.set_alpha(0.0)
            ax = fig.add_subplot(111)
            ax.set_axis_off()
            ax.set_facecolor((1.0, 1.0, 1.0, 0.0))
            ax.text(
                0.02,
                0.5,
                expression,
                fontsize=FORMULA_RENDER_FONT_SIZE,
                va="center",
                ha="left",
                color="#111111",
            )

            FigureCanvasAgg(fig).draw()
            buffer = io.BytesIO()
            fig.savefig(
                buffer,
                format="png",
                dpi=FORMULA_RENDER_DPI,
                transparent=True,
                bbox_inches="tight",
                pad_inches=0.08,
            )
            buffer.seek(0)
            with Image.open(buffer) as rendered:
                rendered.convert("RGBA").save(target_path, format="PNG")
        finally:
            fig.clear()

    def _insert_attachment_token(self, image_id, image_path, surround_with_newlines):
        token = f"[[IMG:{image_id}]]"
        try:
            selection_start = self.text.index("sel.first")
            selection_end = self.text.index("sel.last")
            self.text.delete(selection_start, selection_end)
            insert_at = selection_start
        except tk.TclError:
            insert_at = self.text.index("insert")

        if surround_with_newlines:
            self.text.insert(insert_at, f"\n{token}\n")
            token_start = self.text.index(f"{insert_at}+1c")
            token_end = self.text.index(f"{token_start}+{len(token)}c")
            next_insert = self.text.index(f"{token_end}+2c")
        else:
            self.text.insert(insert_at, token)
            token_start = insert_at
            token_end = self.text.index(f"{token_start}+{len(token)}c")
            next_insert = token_end

        self._render_token_as_inline_image(token_start, token_end, image_path)
        self.text.mark_set("insert", next_insert)
        self.text.see("insert")
        return token

    def insert_formula(self):
        initial_value = ""
        try:
            initial_value = self._normalize_latex_expression(
                self.text.get("sel.first", "sel.last")
            )
        except tk.TclError:
            initial_value = ""

        formula = simpledialog.askstring(
            "Sett inn formel",
            (
                "Skriv LaTeX-formel.\n"
                "Eksempel: \\frac{a+b}{n}, x_i^2, \\sum_{i=1}^{n} i\n"
                "Du kan skrive med eller uten $...$."
            ),
            initialvalue=initial_value or "\\frac{a+b}{n}",
            parent=self.root,
        )
        if formula is None:
            return

        normalized_formula = self._normalize_latex_expression(formula)
        if not normalized_formula:
            self._set_status("Formel avbrutt: tom formel.")
            return

        image_id = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
        image_path = IMAGES_DIR / f"{image_id}.png"
        try:
            self._render_formula_to_image(normalized_formula, image_path)
        except Exception as exc:
            messagebox.showerror(
                "Formel feilet",
                (
                    "Klarte ikke å rendre formel.\n"
                    "Sjekk LaTeX-syntaks (MathText) og prøv igjen.\n\n"
                    f"Feil: {exc}"
                ),
                parent=self.root,
            )
            self._set_status("Formel feilet.")
            return

        self.attachments[image_id] = str(image_path)
        self.formula_meta[image_id] = normalized_formula
        self._insert_attachment_token(
            image_id, str(image_path), surround_with_newlines=False
        )
        self._dirty = True
        self._schedule_spellcheck()
        self._set_status("Formel satt inn (Ctrl+M).")

    def _load_inline_photo(self, image_path):
        try:
            with Image.open(image_path) as img:
                preview = img.convert("RGB")
        except Exception:
            return None

        resampling = getattr(Image, "Resampling", Image)
        preview.thumbnail(
            (INLINE_IMAGE_MAX_WIDTH, INLINE_IMAGE_MAX_HEIGHT), resampling.LANCZOS
        )
        return ImageTk.PhotoImage(preview)

    def _render_token_as_inline_image(self, token_start, token_end, image_path):
        if IMAGE_TOKEN_HIDDEN_TAG in self.text.tag_names(token_start):
            return

        photo = self._load_inline_photo(image_path)
        if photo is None:
            return

        self.text.tag_add(IMAGE_TOKEN_HIDDEN_TAG, token_start, token_end)
        self.text.image_create(token_end, image=photo, padx=3, pady=3)
        self._inline_image_refs.append(photo)

    def _render_inline_images_from_tokens(self):
        search_from = "1.0"
        while True:
            token_start = self.text.search("[[IMG:", search_from, stopindex="end")
            if not token_start:
                return

            token_close = self.text.search("]]", token_start, stopindex="end")
            if not token_close:
                return

            token_end = self.text.index(f"{token_close}+2c")
            token_value = self.text.get(token_start, token_end)
            match = IMAGE_TOKEN_PATTERN.fullmatch(token_value)
            if match and IMAGE_TOKEN_HIDDEN_TAG not in self.text.tag_names(token_start):
                image_id = match.group(1)
                image_path = self.attachments.get(image_id)
                if image_path and Path(image_path).exists():
                    self._render_token_as_inline_image(token_start, token_end, image_path)

            search_from = token_end

    def _save_image_and_insert_token(self, image):
        image_id = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
        image_path = IMAGES_DIR / f"{image_id}.png"
        image.convert("RGB").save(image_path, format="PNG")
        self.attachments[image_id] = str(image_path)
        return self._insert_attachment_token(
            image_id, str(image_path), surround_with_newlines=True
        )

    def _serialize_tag_ranges(self, tag_name):
        ranges = self.text.tag_ranges(tag_name)
        serialized = []
        for i in range(0, len(ranges), 2):
            serialized.append([str(ranges[i]), str(ranges[i + 1])])
        return serialized

    def _restore_tag_ranges(self, tag_name, serialized_ranges):
        for pair in serialized_ranges:
            if not isinstance(pair, list) or len(pair) != 2:
                continue
            try:
                self.text.tag_add(tag_name, pair[0], pair[1])
            except tk.TclError:
                continue

    def _build_session_payload(self):
        return {
            "text": self.text.get("1.0", "end-1c"),
            "attachments": self.attachments,
            "formula_meta": self.formula_meta,
            "bold_ranges": self._serialize_tag_ranges(BOLD_TAG),
            "superscript_ranges": self._serialize_tag_ranges(SUPERSCRIPT_TAG),
        }

    def _write_json_atomic(self, path, payload):
        temp_path = path.with_suffix(path.suffix + ".tmp")
        content = json.dumps(payload, ensure_ascii=False, indent=2)
        temp_path.write_text(content, encoding="utf-8")
        temp_path.replace(path)

    def _read_json_payload(self, path):
        if not path.exists():
            return None
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return None
        return payload if isinstance(payload, dict) else None

    def _payload_has_content(self, payload):
        text = payload.get("text", "")
        if isinstance(text, str) and text.strip():
            return True
        attachments = payload.get("attachments", {})
        return isinstance(attachments, dict) and bool(attachments)

    def _payload_snapshot_key(self, payload):
        return json.dumps(payload, ensure_ascii=False, sort_keys=True)

    def _sorted_backup_files(self):
        return sorted(
            BACKUP_DIR.glob("session-*.json"), key=lambda p: p.stat().st_mtime, reverse=True
        )

    def _prune_backup_files(self):
        backups = self._sorted_backup_files()
        for old_file in backups[MAX_SESSION_BACKUPS:]:
            try:
                old_file.unlink()
            except Exception:
                continue

    def _write_backup_snapshot(self, payload):
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        target = BACKUP_DIR / f"session-{stamp}.json"
        counter = 1
        while target.exists():
            target = BACKUP_DIR / f"session-{stamp}-{counter}.json"
            counter += 1
        self._write_json_atomic(target, payload)
        self._prune_backup_files()

    def _maybe_write_backup_snapshot(self, payload, force=False):
        if not self._payload_has_content(payload):
            return

        snapshot_key = self._payload_snapshot_key(payload)
        now = datetime.now()
        due_by_time = (
            self._last_backup_snapshot_at is None
            or (now - self._last_backup_snapshot_at).total_seconds() * 1000
            >= AUTOSAVE_BACKUP_MS
        )
        changed = snapshot_key != self._last_backup_snapshot_key
        if force or (due_by_time and changed):
            self._write_backup_snapshot(payload)
            self._last_backup_snapshot_at = now
            self._last_backup_snapshot_key = snapshot_key

    def _apply_session_payload(self, payload):
        text_value = payload.get("text", "")
        attachment_map = payload.get("attachments", {})
        formula_map = payload.get("formula_meta", {})
        if not isinstance(attachment_map, dict):
            attachment_map = {}
        if not isinstance(formula_map, dict):
            formula_map = {}
        bold_ranges = payload.get("bold_ranges", [])
        superscript_ranges = payload.get("superscript_ranges", [])

        self.text.delete("1.0", "end")
        self._clear_identical_highlight()
        self._last_identical_token = ""
        self._clear_search_highlights()
        self._inline_image_refs = []
        self.text.insert("1.0", text_value)
        self.attachments = {
            key: value for key, value in attachment_map.items() if Path(value).exists()
        }
        self.formula_meta = {}
        for image_id, raw_formula in formula_map.items():
            if image_id not in self.attachments:
                continue
            if not isinstance(image_id, str):
                continue
            if not isinstance(raw_formula, str):
                continue
            normalized_formula = self._normalize_latex_expression(raw_formula)
            if normalized_formula:
                self.formula_meta[image_id] = normalized_formula
        self._restore_tag_ranges(BOLD_TAG, bold_ranges)
        self._restore_tag_ranges(SUPERSCRIPT_TAG, superscript_ranges)
        self._render_inline_images_from_tokens()
        self.text.edit_modified(False)
        self._dirty = False

    def _save_session(self, silent, force_backup=False):
        payload = self._build_session_payload()
        try:
            if SESSION_FILE.exists():
                try:
                    SESSION_PREV_FILE.write_text(
                        SESSION_FILE.read_text(encoding="utf-8"), encoding="utf-8"
                    )
                except Exception:
                    pass

            self._write_json_atomic(SESSION_FILE, payload)
            self._maybe_write_backup_snapshot(payload, force=force_backup)
        except Exception as exc:
            self._set_status(f"Lagring feilet: {exc}")
            return

        self._dirty = False
        if not silent:
            self._set_status(f"Lagret: {SESSION_FILE}")

    def _load_session(self):
        payload = self._read_json_payload(SESSION_FILE)
        payload_source = "session"

        if payload is None:
            payload = self._read_json_payload(SESSION_PREV_FILE)
            payload_source = "session.prev"

        if payload is None:
            for backup_file in self._sorted_backup_files():
                payload = self._read_json_payload(backup_file)
                if payload is not None:
                    payload_source = backup_file.name
                    break

        if payload is None:
            return

        main_payload = self._read_json_payload(SESSION_FILE)
        if (
            main_payload is not None
            and not self._payload_has_content(main_payload)
            and payload_source == "session"
        ):
            for backup_file in self._sorted_backup_files():
                backup_payload = self._read_json_payload(backup_file)
                if backup_payload and self._payload_has_content(backup_payload):
                    if messagebox.askyesno(
                        "Gjenopprett autosave",
                        (
                            "Fant en autosave-kopi med innhold.\n"
                            f"Vil du laste den i stedet?\n\n{backup_file.name}"
                        ),
                    ):
                        payload = backup_payload
                        payload_source = backup_file.name
                    break

        self._apply_session_payload(payload)

        if self._payload_has_content(payload):
            if payload_source == "session":
                self._set_status(
                    f"Gjenopprettet forrige notat ({len(self.attachments)} bilder koblet til)."
                )
            else:
                self._set_status(
                    f"Gjenopprettet notat fra backup ({payload_source})."
                )

        self._schedule_spellcheck()

    def _split_text_runs_by_format(self, segment_start, segment_text):
        if not segment_text:
            return []

        runs = []
        current_style = None
        buffer = []

        for offset, ch in enumerate(segment_text):
            char_index = self.text.index(f"{segment_start}+{offset}c")
            is_bold = BOLD_TAG in self.text.tag_names(char_index)
            is_superscript = SUPERSCRIPT_TAG in self.text.tag_names(char_index)
            style = (is_bold, is_superscript)
            if current_style is None:
                current_style = style

            if style != current_style:
                runs.append(("".join(buffer), current_style[0], current_style[1]))
                buffer = []
                current_style = style

            buffer.append(ch)

        if buffer:
            runs.append(("".join(buffer), current_style[0], current_style[1]))

        return runs

    def _runs_to_paragraphs(self, runs):
        paragraphs = [[]]
        for run_text, is_bold, is_superscript in runs:
            parts = run_text.split("\n")
            for i, part in enumerate(parts):
                if part:
                    paragraphs[-1].append((part, is_bold, is_superscript))
                if i < len(parts) - 1:
                    paragraphs.append([])

        if not paragraphs:
            return [[]]
        return paragraphs

    def _normalize_paragraph_runs_for_export(self, paragraph_runs):
        if not paragraph_runs:
            return paragraph_runs

        chars = []
        for run_text, is_bold, is_superscript in paragraph_runs:
            for ch in run_text:
                chars.append([ch, is_bold, is_superscript])

        if not chars:
            return paragraph_runs

        plain_text = "".join(item[0] for item in chars)
        for match in re.finditer(r"[A-Za-z\u00C0-\u024F]+", plain_text):
            start, end = match.span()
            bold_flags = [chars[i][1] for i in range(start, end)]
            if not any(bold_flags) or all(bold_flags):
                continue

            prefix_nonbold = 0
            while prefix_nonbold < len(bold_flags) and not bold_flags[prefix_nonbold]:
                prefix_nonbold += 1

            if 0 < prefix_nonbold <= 2 and all(bold_flags[prefix_nonbold:]):
                for i in range(start, start + prefix_nonbold):
                    chars[i][1] = True

        normalized = []
        buffer = []
        current_style = None

        for ch, is_bold, is_superscript in chars:
            style = (is_bold, is_superscript)
            if current_style is None:
                current_style = style

            if style != current_style:
                normalized.append(("".join(buffer), current_style[0], current_style[1]))
                buffer = [ch]
                current_style = style
            else:
                buffer.append(ch)

        if buffer:
            normalized.append(("".join(buffer), current_style[0], current_style[1]))

        return normalized

    def _content_blocks_for_export(self):
        text_value = self.text.get("1.0", "end-1c")
        blocks = []
        cursor = 0

        for match in IMAGE_TOKEN_PATTERN.finditer(text_value):
            if match.start() > cursor:
                segment_text = text_value[cursor : match.start()]
                segment_start = self.text.index(f"1.0+{cursor}c")
                blocks.append(
                    ("text_runs", self._split_text_runs_by_format(segment_start, segment_text))
                )

            image_id = match.group(1)
            image_path = self.attachments.get(image_id)
            if image_path and Path(image_path).exists():
                blocks.append(("image", image_path))
            else:
                blocks.append(("text_runs", [(f"[Mangler bilde: {image_id}]", False, False)]))

            cursor = match.end()

        if cursor < len(text_value):
            segment_text = text_value[cursor:]
            segment_start = self.text.index(f"1.0+{cursor}c")
            blocks.append(
                ("text_runs", self._split_text_runs_by_format(segment_start, segment_text))
            )

        if not blocks:
            blocks.append(("text_runs", [("", False, False)]))

        return blocks

    def _trim_prefix_from_runs(self, runs, prefix_len):
        remaining = max(0, int(prefix_len))
        trimmed = []
        for run_text, is_bold, is_superscript in runs:
            if remaining >= len(run_text):
                remaining -= len(run_text)
                continue
            if remaining > 0:
                run_text = run_text[remaining:]
                remaining = 0
            if run_text:
                trimmed.append((run_text, is_bold, is_superscript))
        return trimmed

    def _heading_level_from_runs(self, runs, plain_text):
        text = plain_text.strip()
        if not text:
            return 0

        numbered_match = re.match(r"^\s*(\d+(?:\.\d+)*)[.)]?\s+\S", plain_text)
        if numbered_match:
            return min(3, numbered_match.group(1).count(".") + 1)

        if len(text) > 110:
            return 0

        total_letters = 0
        bold_letters = 0
        for run_text, is_bold, _is_superscript in runs:
            for ch in run_text:
                if not ch.isalpha():
                    continue
                total_letters += 1
                if is_bold:
                    bold_letters += 1

        if total_letters >= 3 and bold_letters / max(total_letters, 1) >= 0.8:
            return 1
        return 0

    def _classify_paragraph_for_export(self, paragraph_runs):
        plain_text = "".join(text for text, _is_bold, _is_superscript in paragraph_runs)
        stripped = plain_text.strip()
        if not stripped:
            return {"kind": "blank", "runs": []}

        bullet_info = self._parse_bullet_line(plain_text)
        if bullet_info:
            content_runs = self._trim_prefix_from_runs(paragraph_runs, bullet_info["prefix_len"])
            return {
                "kind": "bullet",
                "level": int(bullet_info["level"]),
                "marker": bullet_info["marker"],
                "runs": content_runs,
            }

        heading_level = self._heading_level_from_runs(paragraph_runs, plain_text)
        if heading_level > 0:
            return {
                "kind": "heading",
                "level": int(heading_level),
                "runs": paragraph_runs,
            }

        return {"kind": "normal", "runs": paragraph_runs}

    def export_word(self):
        default_name = f"notater-{datetime.now():%Y%m%d-%H%M}.docx"
        target = filedialog.asksaveasfilename(
            title="Eksporter til Word",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word-dokument", "*.docx")],
        )
        if not target:
            return

        doc = Document()
        for block_type, value in self._content_blocks_for_export():
            if block_type == "text_runs":
                for paragraph_runs in self._runs_to_paragraphs(value):
                    paragraph_runs = self._normalize_paragraph_runs_for_export(paragraph_runs)
                    paragraph_info = self._classify_paragraph_for_export(paragraph_runs)
                    kind = paragraph_info["kind"]
                    runs = paragraph_info["runs"]
                    marker_prefix = ""

                    if kind == "blank":
                        doc.add_paragraph()
                        continue

                    if kind == "heading":
                        style_name = f"Heading {min(3, paragraph_info['level'])}"
                        try:
                            paragraph = doc.add_paragraph(style=style_name)
                        except Exception:
                            paragraph = doc.add_paragraph()
                    elif kind == "bullet":
                        level = max(0, int(paragraph_info["level"]))
                        style_name = "List Bullet" if level == 0 else f"List Bullet {min(3, level + 1)}"
                        try:
                            paragraph = doc.add_paragraph(style=style_name)
                        except Exception:
                            paragraph = doc.add_paragraph()
                            marker_prefix = f"{paragraph_info['marker']} "
                    else:
                        paragraph = doc.add_paragraph()

                    if marker_prefix:
                        paragraph.add_run(marker_prefix)

                    for run_text, is_bold, is_superscript in runs:
                        run = paragraph.add_run(run_text)
                        run.bold = bool(is_bold)
                        run.font.superscript = bool(is_superscript)
            else:
                try:
                    doc.add_picture(value, width=Inches(5.7))
                except Exception:
                    doc.add_paragraph(f"[Kunne ikke legge til bilde: {Path(value).name}]")

        doc.save(target)
        self._set_status(f"Eksporterte Word: {target}")
        messagebox.showinfo("Eksport fullført", f"Lagret Word-fil:\n{target}")

    def export_pdf(self):
        default_name = f"notater-{datetime.now():%Y%m%d-%H%M}.pdf"
        target = filedialog.asksaveasfilename(
            title="Eksporter til PDF",
            defaultextension=".pdf",
            initialfile=default_name,
            filetypes=[("PDF", "*.pdf")],
        )
        if not target:
            return

        doc = SimpleDocTemplate(
            target,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36,
        )
        styles = getSampleStyleSheet()
        normal = styles["Normal"].clone("NotesNormal")
        normal.leading = max(normal.leading, int(normal.fontSize * self.pdf_line_spacing))

        heading_styles = {}
        for level in (1, 2, 3):
            style = normal.clone(f"NotesHeading{level}")
            style.fontName = "Helvetica-Bold"
            style.fontSize = max(11, 16 - (level - 1) * 2)
            style.leading = max(style.leading, int(style.fontSize * 1.35))
            style.spaceBefore = 10 if level == 1 else 7
            style.spaceAfter = 4
            heading_styles[level] = style

        bullet_styles = {}
        for level in range(0, 6):
            style = normal.clone(f"NotesBullet{level}")
            indent = 14 + level * 16
            style.leftIndent = indent
            style.firstLineIndent = 0
            style.bulletIndent = max(0, indent - 10)
            style.spaceBefore = 1
            style.spaceAfter = 1
            bullet_styles[level] = style
        story = []

        for block_type, value in self._content_blocks_for_export():
            if block_type == "text_runs":
                for paragraph_runs in self._runs_to_paragraphs(value):
                    paragraph_runs = self._normalize_paragraph_runs_for_export(paragraph_runs)
                    paragraph_info = self._classify_paragraph_for_export(paragraph_runs)
                    kind = paragraph_info["kind"]
                    runs = paragraph_info["runs"]
                    if kind == "blank":
                        story.append(Paragraph("&nbsp;", normal))
                        continue

                    paragraph_chunks = []
                    for run_text, is_bold, is_superscript in runs:
                        escaped_text = escape(run_text)
                        chunk = escaped_text
                        if is_superscript:
                            chunk = f"<super>{chunk}</super>"
                        if is_bold:
                            chunk = f"<b>{chunk}</b>"
                        paragraph_chunks.append(chunk)

                    paragraph_markup = "".join(paragraph_chunks) or "&nbsp;"
                    if kind == "heading":
                        level = min(3, int(paragraph_info["level"]))
                        story.append(Paragraph(paragraph_markup, heading_styles[level]))
                    elif kind == "bullet":
                        level = max(0, min(5, int(paragraph_info["level"])))
                        marker = paragraph_info.get("marker") or "\u2022"
                        story.append(
                            Paragraph(
                                paragraph_markup,
                                bullet_styles[level],
                                bulletText=marker,
                            )
                        )
                    else:
                        story.append(Paragraph(paragraph_markup, normal))
                story.append(Spacer(1, 8))
            else:
                try:
                    with Image.open(value) as img:
                        width, height = img.size
                    max_width = 520
                    if width > max_width:
                        scale = max_width / width
                        width = width * scale
                        height = height * scale
                    story.append(RLImage(value, width=width, height=height))
                    story.append(Spacer(1, 10))
                except Exception:
                    story.append(
                        Paragraph(
                            escape(f"[Kunne ikke legge til bilde: {Path(value).name}]"),
                            normal,
                        )
                    )

        doc.build(story)
        self._set_status(f"Eksporterte PDF: {target}")
        messagebox.showinfo("Eksport fullført", f"Lagret PDF-fil:\n{target}")

    def on_close(self):
        if self._fade_after_id is not None:
            self.root.after_cancel(self._fade_after_id)
            self._fade_after_id = None
        if self._spellcheck_after_id is not None:
            self.root.after_cancel(self._spellcheck_after_id)
            self._spellcheck_after_id = None
        if self._toc_window is not None and self._toc_window.winfo_exists():
            self._toc_window.destroy()
            self._toc_window = None
        if self._settings_window is not None and self._settings_window.winfo_exists():
            self._settings_window.destroy()
            self._settings_window = None
        self._save_session(silent=True, force_backup=True)
        self.hotkey.stop()
        self.root.destroy()

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = NoteOverlayApp()
    app.run()
